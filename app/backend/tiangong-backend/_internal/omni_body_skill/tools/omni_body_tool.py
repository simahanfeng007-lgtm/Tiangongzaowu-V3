"""
Tiangong Omni Body Tool / App Bus
=================================

A single-entry v3 executable tool and application-capability bus.

Important design rule:
- This is NOT a vague universal_tool(goal) black box.
- The model must call body.run(action=..., target=..., args=...).
- Each action returns evidence suitable for the next reasoning step.
- Mutating actions snapshot affected paths first, so rollback.apply can restore state.

Portable core actions are implemented in pure Python + common libraries.
Actions requiring external GUI/cloud/model backends return a structured
"requires_adapter" response instead of pretending to work.
"""

from __future__ import annotations

import base64
from contextlib import contextmanager
import csv
import hashlib
import html as html_lib
import json
import math
import mimetypes
import os
import re
import shlex
import shutil
import struct
import subprocess
import sys
import tempfile
import threading
import time
import traceback
import uuid
import zipfile
import urllib.parse
import urllib.request
import unicodedata
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
from ._stub_actions import _stub_action_result
from .sandbox_runtime import (
    WINDOWS_UTF8_SHELL_MARKER,
    SandboxLimits,
    SandboxRunner,
    _prepare_windows_utf8_shell_command,
    _snapshot as _sandbox_workspace_snapshot,
    sanitized_environment,
)
from .portable_text import PortableTextError, decode_portable_bytes, subprocess_environment

# The bundled skill can be loaded either through backend.v3 or as a standalone
# `tools` package. Locate the backend root before importing the shared fact kernel.
try:
    from v3.fact_kernel import FactExecutionKernel, compile_manifest, fact_execution_active
except Exception:
    for _parent in Path(__file__).resolve().parents:
        if _parent.name == "v3" and (_parent / "fact_kernel").is_dir():
            _backend_root = str(_parent.parent)
            if _backend_root not in sys.path:
                sys.path.insert(0, _backend_root)
            break
    from v3.fact_kernel import FactExecutionKernel, compile_manifest, fact_execution_active

try:
    from html.parser import HTMLParser as _HTMLParser
except ModuleNotFoundError:
    _HTMLParser = None  # type: ignore[assignment]

try:
    import wave
except ModuleNotFoundError:
    wave = None  # type: ignore[assignment]

try:
    from ..tool_contracts import schema_for_action, validate_tool_request
except ImportError:  # standalone ``tools`` package used by the bundled test suite
    from tool_contracts import schema_for_action, validate_tool_request


# D-21 用户指定路径根放行（与 tool_contracts._user_specified_allowed 同一口径；
# 本文件自含一份以便 AST 抽取/单文件加载路径 fail-closed 兜底）。
def _user_path_root_allowed(resolved_text: str, user_roots) -> bool:
    if not user_roots:
        return False
    resolved = Path(str(resolved_text))
    text = os.path.normcase(str(resolved))
    for prefix in (
        os.path.normcase(os.environ.get("SystemRoot") or r"C:\Windows"),
        os.path.normcase(os.environ.get("ProgramFiles") or r"C:\Program Files"),
        os.path.normcase(os.environ.get("ProgramFiles(x86)") or r"C:\Program Files (x86)"),
        os.path.normcase(os.environ.get("ProgramData") or r"C:\ProgramData"),
    ):
        if text == prefix or text.startswith(prefix + os.sep):
            return False
    if len(text) <= 3:
        return False
    if any(part.casefold() in {".ssh", ".aws", ".gnupg", ".azure", ".config"} for part in resolved.parts):
        return False
    if resolved.name.casefold().endswith(".env"):
        return False
    for root in user_roots:
        try:
            resolved.relative_to(Path(str(root)).resolve(strict=False))
            return True
        except (ValueError, OSError):
            continue
    return False


_WORKSPACE_LOCKS_GUARD = threading.Lock()
_WORKSPACE_LOCKS: Dict[str, threading.RLock] = {}
_WORKSPACE_LOCK_STATE = threading.local()
_LIFE_ACTIVITY_QUERY_PROVIDER: Any = None
_BODY_STATE_QUERY_PROVIDER: Any = None
_LEARNING_INGEST_PROVIDER: Any = None


def set_life_activity_query_provider(provider: Any) -> None:
    """Inject the authoritative in-process Life activity reader."""

    global _LIFE_ACTIVITY_QUERY_PROVIDER
    if provider is not None and not callable(provider):
        raise TypeError("life activity query provider must be callable")
    _LIFE_ACTIVITY_QUERY_PROVIDER = provider


def set_body_state_query_provider(provider: Any) -> None:
    """Inject the Total Gateway-owned, read-only body state projection."""

    global _BODY_STATE_QUERY_PROVIDER
    if provider is not None and not callable(provider):
        raise TypeError("body state query provider must be callable")
    _BODY_STATE_QUERY_PROVIDER = provider


def set_learning_ingest_provider(provider: Any) -> None:
    """Inject the authoritative pending-only Life learning writer."""

    global _LEARNING_INGEST_PROVIDER
    if provider is not None and not callable(provider):
        raise TypeError("learning ingest provider must be callable")
    _LEARNING_INGEST_PROVIDER = provider


def _workspace_lock_for(path: Path) -> threading.RLock:
    key = str(path.resolve())
    with _WORKSPACE_LOCKS_GUARD:
        lock = _WORKSPACE_LOCKS.get(key)
        if lock is None:
            lock = threading.RLock()
            _WORKSPACE_LOCKS[key] = lock
        return lock




def _workspace_lock_timeout_seconds() -> float:
    """Parse the lock timeout without letting a malformed env var stop work."""
    raw = str(os.environ.get("TIANGONG_WORKSPACE_LOCK_TIMEOUT_SECONDS") or "").strip()
    try:
        parsed = float(raw) if raw else 900.0
    except (TypeError, ValueError, OverflowError):
        parsed = 900.0
    if not math.isfinite(parsed):
        parsed = 900.0
    return max(10.0, min(parsed, 3600.0))

def _lock_file_handle(handle: Any) -> None:
    """Acquire a bounded cross-process lock on the first byte of *handle*."""
    if os.name == "nt":
        import msvcrt

        handle.seek(0, os.SEEK_END)
        if handle.tell() == 0:
            handle.write(b"\0")
            handle.flush()
        timeout = _workspace_lock_timeout_seconds()
        deadline = time.monotonic() + timeout
        while True:
            handle.seek(0)
            try:
                msvcrt.locking(handle.fileno(), msvcrt.LK_NBLCK, 1)
                return
            except OSError:
                if time.monotonic() >= deadline:
                    raise TimeoutError("workspace_cross_process_lock_timeout")
                time.sleep(0.02)

    import fcntl

    timeout = _workspace_lock_timeout_seconds()
    deadline = time.monotonic() + timeout
    while True:
        try:
            fcntl.flock(handle.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
            return
        except BlockingIOError:
            if time.monotonic() >= deadline:
                raise TimeoutError("workspace_cross_process_lock_timeout")
            time.sleep(0.02)


def _unlock_file_handle(handle: Any) -> None:
    if os.name == "nt":
        import msvcrt

        handle.seek(0)
        msvcrt.locking(handle.fileno(), msvcrt.LK_UNLCK, 1)
        return

    import fcntl

    fcntl.flock(handle.fileno(), fcntl.LOCK_UN)


@contextmanager
def _workspace_mutation_guard(path: Path, thread_lock: threading.RLock):
    """Serialize a workspace mutation across threads and OS processes.

    The per-thread depth map makes nested alias actions and audit writes
    re-entrant while a single outer file lock protects independent Agent
    processes sharing the same workspace.
    """
    key = str(path.resolve())
    with thread_lock:
        depths = getattr(_WORKSPACE_LOCK_STATE, "depths", None)
        handles = getattr(_WORKSPACE_LOCK_STATE, "handles", None)
        if depths is None:
            depths = {}
            handles = {}
            _WORKSPACE_LOCK_STATE.depths = depths
            _WORKSPACE_LOCK_STATE.handles = handles

        depth = int(depths.get(key, 0))
        handle = None
        if depth == 0:
            path.mkdir(parents=True, exist_ok=True)
            lock_path = path / ".omni_workspace.lock"
            handle = lock_path.open("a+b")
            try:
                _lock_file_handle(handle)
            except Exception:
                handle.close()
                raise
            handles[key] = handle
        depths[key] = depth + 1
        try:
            yield
        finally:
            remaining = int(depths.get(key, 1)) - 1
            if remaining > 0:
                depths[key] = remaining
            else:
                depths.pop(key, None)
                held = handles.pop(key, handle)
                if held is not None:
                    try:
                        _unlock_file_handle(held)
                    finally:
                        held.close()


RISK = {
    "A0": "read-only / observation",
    "A1": "local temporary creation",
    "A2": "create durable local output",
    "A3": "modify or move user files with rollback snapshot",
    "A4": "delete-to-trash, shell, extraction overwrite, external write; automatic in sandbox; A5 is hard rejected",
    "A5": "hard blocked: payment, credential theft, bypass, destructive permanent delete, unauthorized voice cloning",
}

ACTIONS: Dict[str, Dict[str, Any]] = {
    "life.body.state.query": {
        "risk": "A0",
        "implemented": True,
        "effect": "read",
        "summary": "Read this Life's current authoritative Life projection and live runtime body state with a gateway audit receipt.",
    },
    "life.activity.query": {
        "risk": "A0",
        "implemented": True,
        "summary": "Read the authoritative Life activity ledger for today, yesterday, or a specified date.",
    },
    "system.capabilities": {"risk": "A0", "implemented": True, "summary": "List supported actions, schemas, risk levels, and adapter-only capabilities."},
    "system.health": {"risk": "A0", "implemented": True, "summary": "Inspect workspace, dependencies, ffmpeg availability, and runtime configuration."},
    "system.app_registry": {"risk": "A0", "implemented": True, "summary": "List mounted application tool groups and their action names."},
    "system.action_schema": {"risk": "A0", "implemented": True, "summary": "Return metadata for a single action, including aliases and adapter requirements."},
    "learning.ingest": {
        "risk": "A2",
        "implemented": True,
        "summary": "Create a pending learning card from a host-verified explicit user learning request; never compiles, activates, registers, or releases tools.",
        "allowed_effect": "create_pending_learning_card_only",
        "requires_explicit_user_learning_intent": True,
        "requires_host_verified_intent_token": True,
    },

    "file.list": {"risk": "A0", "implemented": True, "summary": "List files/directories under a target directory."},
    "file.read": {"risk": "A0", "implemented": True, "summary": "Read text or base64-encoded binary file preview."},
    "file.write": {"risk": "A3", "implemented": True, "summary": "Write text/base64 file with rollback snapshot."},
    "file.append": {"risk": "A3", "implemented": True, "summary": "Append text to a file with rollback snapshot."},
    "file.copy": {"risk": "A3", "implemented": True, "summary": "Copy a file or directory with overwrite protection and rollback snapshot."},
    "file.move": {"risk": "A3", "implemented": True, "summary": "Move a file or directory with rollback snapshot."},
    "file.rename": {"risk": "A3", "implemented": True, "summary": "Rename a file or directory in place with rollback snapshot."},
    "file.mkdir": {"risk": "A2", "implemented": True, "summary": "Create a directory."},
    "file.delete_to_trash": {"risk": "A4", "implemented": True, "summary": "Move file/directory into workspace trash; never permanent delete."},
    "file.search": {"risk": "A0", "implemented": True, "summary": "Search filenames and/or text content within workspace."},
    "file.hash": {"risk": "A0", "implemented": True, "summary": "Calculate SHA-256 for a file."},

    "zip.create": {"risk": "A2", "implemented": True, "summary": "Create a zip archive from files/directories."},
    "zip.extract": {"risk": "A4", "implemented": True, "summary": "Safely extract zip archive under destination with zip-slip checks."},

    "code.read": {"risk": "A0", "implemented": True, "summary": "Read source code file with metadata."},
    "code.write": {"risk": "A3", "implemented": True, "summary": "Write source code and optionally run syntax checks."},
    "code.patch_replace": {"risk": "A3", "implemented": True, "summary": "Patch file by replacing literal or regex text with rollback snapshot."},
    "file.patch_replace": {"risk": "A3", "implemented": True, "alias_to": "code.patch_replace", "summary": "Compatibility alias for code.patch_replace."},
    "quality.python_syntax": {"risk": "A0", "implemented": True, "summary": "Compile Python files to detect syntax errors."},
    "quality.javascript_syntax": {"risk": "A0", "implemented": True, "summary": "Run node --check on JavaScript files without a free-form shell."},
    "quality.run_tests": {"risk": "A4", "implemented": True, "summary": "Run a test command inside workspace; requires allow_shell=True."},
    "python.run": {"risk": "A4", "implemented": True, "summary": "Run Python code or a Python script in workspace; requires allow_python=True."},
    "shell.run": {"risk": "A4", "implemented": True, "summary": "Run shell command in workspace; requires allow_shell=True."},
    "git.clone": {"risk": "A2", "implemented": True, "effect": "create", "summary": "Clone a public GitHub HTTPS repository into a new local directory through a governed network-read capability; generic shell/python remain network-denied."},

    "docx.create": {"risk": "A2", "implemented": True, "summary": "Create a Word .docx document from structured sections/tables."},
    "word.create": {"risk": "A2", "implemented": True, "summary": "Alias for docx.create — create a Word document."},
    "pptx.create": {"risk": "A2", "implemented": True, "summary": "Create a PowerPoint .pptx deck from structured slide specs."},
    "pptx.read": {"risk": "A0", "implemented": True, "summary": "Inspect PowerPoint slide text, dimensions, placeholders, fonts, and meaningful visual evidence."},
    "sheet.create": {"risk": "A2", "implemented": True, "summary": "Create an .xlsx workbook from structured sheets."},
    "sheet.read": {"risk": "A0", "implemented": True, "summary": "Read preview rows from .xlsx or .csv."},

    # —— 别名映射：word.* → docx.* ——
    "word.read": {"risk": "A0", "implemented": True, "summary": "Alias for docx: use file.read for .docx files."},
    "mindmap.create": {"risk": "A2", "implemented": True, "summary": "Create Mermaid mindmap markdown and optional OPML."},

    "pdf.extract_text": {"risk": "A0", "implemented": True, "summary": "Extract text from a PDF via pypdf."},
    "pdf.create_from_text": {"risk": "A2", "implemented": True, "summary": "Create a simple PDF from text via reportlab."},

    "image.info": {"risk": "A0", "implemented": True, "summary": "Inspect image size/mode/format."},
    "image.create_canvas": {"risk": "A2", "implemented": True, "summary": "Create a blank image canvas."},
    "image.resize": {"risk": "A2", "implemented": True, "summary": "Resize image with Pillow."},
    "image.crop": {"risk": "A2", "implemented": True, "summary": "Crop image with Pillow."},
    "image.rotate": {"risk": "A2", "implemented": True, "summary": "Rotate image with Pillow."},
    "image.add_text": {"risk": "A2", "implemented": True, "summary": "Add simple text overlay to image."},
    "image.compose": {"risk": "A2", "implemented": True, "summary": "Overlay one image on another."},
    "image.convert": {"risk": "A2", "implemented": True, "summary": "Convert image format."},

    "audio.tone": {"risk": "A2", "implemented": True, "summary": "Create a simple WAV tone/beep/music bed placeholder."},
    "audio.trim": {"risk": "A2", "implemented": True, "summary": "Trim audio using ffmpeg."},
    "audio.concat": {"risk": "A2", "implemented": True, "summary": "Concatenate audio files using ffmpeg concat demuxer."},
    "audio.tts": {"risk": "A2", "implemented": False, "adapter": "tts_backend", "summary": "Text-to-speech through configured TTS backend adapter."},
    "voice.clone_authorized": {"risk": "A5", "implemented": False, "adapter": "voice_backend_with_consent", "summary": "Only for owned/consented voices; disabled by default and requires explicit external consent gate."},

    "video.info": {"risk": "A0", "implemented": True, "summary": "Inspect video metadata using ffprobe."},
    "video.cut": {"risk": "A2", "implemented": True, "summary": "Cut a video segment using ffmpeg."},
    "video.extract_audio": {"risk": "A2", "implemented": True, "summary": "Extract audio track using ffmpeg."},
    "video.add_audio": {"risk": "A2", "implemented": True, "summary": "Mux/replace audio track using ffmpeg."},
    "video.slideshow": {"risk": "A2", "implemented": True, "summary": "Create simple slideshow video from images using ffmpeg."},

    "rollback.list": {"risk": "A0", "implemented": True, "summary": "List rollback-capable operations."},
    "rollback.apply": {"risk": "A4", "implemented": True, "summary": "Restore snapshots for a previous mutating operation."},

    "browser.open": {"risk": "A3", "implemented": False, "adapter": "browser_driver", "summary": "Open or operate browser through Playwright/Selenium/GUI adapter."},
    "browser.search_web": {"risk": "A0", "implemented": False, "adapter": "web_search", "summary": "Use host model/web-search tool; this portable package has no bundled web access."},
    "desktop.screenshot": {"risk": "A0", "implemented": False, "adapter": "desktop_automation", "summary": "Screenshot through pyautogui/accessibility adapter."},
    "desktop.click": {"risk": "A3", "implemented": False, "adapter": "desktop_automation", "summary": "Click through GUI adapter."},
    "desktop.type": {"risk": "A3", "implemented": False, "adapter": "desktop_automation", "summary": "Type through GUI adapter."},
    "desktop.hotkey": {"risk": "A3", "implemented": False, "adapter": "desktop_automation", "summary": "Keyboard shortcut through GUI adapter."},
}


# ---------- app-bus registry loader ----------

def _find_registry_root() -> Path | None:
    here = Path(__file__).resolve()
    candidates = []
    for parent in [here.parent, *here.parents]:
        candidates.append(parent)
        candidates.append(parent.parent if parent.parent != parent else parent)
    for root in candidates:
        if (root / "registry" / "app_actions.json").exists():
            return root
    return None


def _load_appbus_registry() -> tuple[dict[str, Any], dict[str, dict[str, Any]]]:
    root = _find_registry_root()
    if root is None:
        return ({"schema": "tiangong.v3.omni_body.app_registry.v1", "apps": []}, {})
    try:
        apps_path = root / "registry" / "apps.json"
        actions_path = root / "registry" / "app_actions.json"
        apps_doc = json.loads(apps_path.read_text(encoding="utf-8")) if apps_path.exists() else {"apps": []}
        actions_doc = json.loads(actions_path.read_text(encoding="utf-8")) if actions_path.exists() else {"actions": {}}
        raw_actions = actions_doc.get("actions", {}) if isinstance(actions_doc, dict) else {}
        actions: dict[str, dict[str, Any]] = {}
        for name, meta in raw_actions.items():
            if isinstance(name, str) and isinstance(meta, dict):
                cleaned = dict(meta)
                cleaned.setdefault("risk", "A2")
                cleaned.setdefault("implemented", False)
                cleaned.setdefault("summary", f"Mounted app-bus action: {name}")
                actions[name] = cleaned
        return (apps_doc if isinstance(apps_doc, dict) else {"apps": []}, actions)
    except Exception as exc:
        return ({"schema": "tiangong.v3.omni_body.app_registry.v1", "apps": [], "load_error": str(exc)}, {})


def _load_learning_runtime() -> tuple[Any | None, Any | None, str]:
    # The legacy autonomous-learning runtime is intentionally detached.  New
    # learning is owned by the embedded Life kernel and reaches the model only
    # through Total Gateway's audited decision route.
    return None, None, "legacy_learning_runtime_detached"


APP_REGISTRY, APP_ACTIONS = _load_appbus_registry()
# Merge mounted application actions into the single tool action table.
# These are NOT registered as separate v3 tools; they are routable actions under omni_body.
ACTIONS.update(APP_ACTIONS)

# Portable app fallbacks: these make the most common app-bus actions truly executable
# without requiring the proprietary application to be installed. They do NOT claim
# to be full native adapters. Native GUI/API adapters can still override these names
# in a host deployment by registering higher-priority handlers before import.
PORTABLE_APP_FALLBACKS: Dict[str, Dict[str, Any]] = {
    # Browser: static fetch/download/text extraction. JS-heavy apps still need Playwright/CDP.
    "browser.open": {"risk": "A2", "implemented": True, "summary": "Portable browser open/fetch fallback: downloads URL/file/data into a local snapshot."},
    "browser.search_web": {"risk": "A0", "implemented": True, "summary": "Portable web search fallback through a configured search endpoint or direct search URL snapshot."},
    "web.search": {"risk": "A0", "implemented": True, "alias_to": "browser.search_web", "summary": "Alias to browser.search_web."},
    "search_web": {"risk": "A0", "implemented": True, "alias_to": "browser.search_web", "summary": "Alias to browser.search_web."},
    "web.read": {"risk": "A0", "implemented": True, "alias_to": "browser.open", "summary": "Read visible text from a URL."},
    "web.fetch": {"risk": "A0", "implemented": True, "alias_to": "browser.open", "summary": "Compatibility alias to web.read."},
    "web_readability_extract": {"risk": "A0", "implemented": True, "alias_to": "browser.open", "summary": "Compatibility alias to web.read."},
    "web.readability_extract": {"risk": "A0", "implemented": True, "alias_to": "browser.open", "summary": "Compatibility alias to web.read."},
    "http.get": {"risk": "A0", "implemented": True, "summary": "Portable HTTP GET fallback returning status, content type, and decoded body text."},
    "web.download": {"risk": "A2", "implemented": True, "summary": "Portable URL download fallback routed to browser.chrome.download."},
    "browser.chrome.open": {"risk": "A2", "implemented": True, "summary": "Portable Chrome-open fallback routed to browser.chrome.goto."},
    "browser.chrome.goto": {"risk": "A2", "implemented": True, "summary": "Portable Chrome goto fallback: fetch URL/file/data and save local HTML/text snapshot."},
    "browser.chrome.extract_text": {"risk": "A0", "implemented": True, "summary": "Extract visible text from a fetched URL or local HTML snapshot."},
    "browser.chrome.extract_dom": {"risk": "A0", "implemented": True, "summary": "Return static DOM/HTML snapshot. Dynamic JS pages need a real browser adapter."},
    "browser.chrome.download": {"risk": "A2", "implemented": True, "summary": "Download a URL or copy file:// resource into workspace."},
    "browser.chrome.pdf.print": {"risk": "A2", "implemented": True, "summary": "Create a simple PDF from static page text. Pixel-perfect print needs browser adapter."},
    "browser.chrome.screenshot": {"risk": "A2", "implemented": True, "summary": "Create a portable text-image screenshot from static page text. Real viewport screenshot needs browser adapter."},

    # Photoshop: portable layer project using JSON + PNG composite. Native PSD/UXP still needs Photoshop adapter.
    "adobe.photoshop.document.create": {"risk": "A2", "implemented": True, "summary": "Create portable Photoshop-like design project with layer JSON and PNG composite."},
    "adobe.photoshop.document.open": {"risk": "A0", "implemented": True, "summary": "Open/read portable Photoshop-like design project metadata."},
    "adobe.photoshop.layer.create": {"risk": "A2", "implemented": True, "summary": "Add a layer to portable Photoshop-like project and re-render PNG composite."},
    "adobe.photoshop.layer.update": {"risk": "A2", "implemented": True, "summary": "Update a portable Photoshop-like project layer and re-render PNG composite."},
    "adobe.photoshop.text.add": {"risk": "A2", "implemented": True, "summary": "Add text layer to portable Photoshop-like project and re-render PNG composite."},
    "adobe.photoshop.export.png": {"risk": "A2", "implemented": True, "summary": "Export portable Photoshop-like project composite to PNG."},
    "adobe.photoshop.image.resize": {"risk": "A2", "implemented": True, "summary": "Resize image through portable Pillow fallback."},
    "adobe.photoshop.image.crop": {"risk": "A2", "implemented": True, "summary": "Crop image through portable Pillow fallback."},

    # Jianying/CapCut: portable video-project JSON + ffmpeg renderer. Native template/effects still need app adapter.
    "jianying.project.create": {"risk": "A2", "implemented": True, "summary": "Create portable Jianying/CapCut-like project JSON."},
    "jianying.media.import": {"risk": "A2", "implemented": True, "summary": "Import media path into portable Jianying project JSON."},
    "jianying.timeline.cut": {"risk": "A2", "implemented": True, "summary": "Add cut segment instruction to portable Jianying project JSON."},
    "jianying.subtitle.add": {"risk": "A2", "implemented": True, "summary": "Add subtitle instruction to portable Jianying project JSON."},
    "jianying.music.add": {"risk": "A2", "implemented": True, "summary": "Attach music/audio path to portable Jianying project JSON."},
    "jianying.cover.create": {"risk": "A2", "implemented": True, "summary": "Create simple cover image for portable Jianying project."},
    "jianying.export.mp4": {"risk": "A2", "implemented": True, "summary": "Render portable Jianying project/video/images to MP4 through ffmpeg."},

    # Feishu: real API can be mounted by env adapter; portable fallback creates local docx/md/pdf deliverables.
    "feishu.docs.doc.create": {"risk": "A2", "implemented": True, "summary": "Create Feishu-like document locally, or remote Feishu doc if configured by host adapter."},
    "feishu.docs.doc.read": {"risk": "A0", "implemented": True, "summary": "Read local Feishu-like document fallback file."},
    "feishu.docs.doc.update": {"risk": "A3", "implemented": True, "summary": "Update local Feishu-like document fallback file."},
    "feishu.docs.export.docx": {"risk": "A2", "implemented": True, "summary": "Export local Feishu-like document fallback to docx."},
    "feishu.docs.export.pdf": {"risk": "A2", "implemented": True, "summary": "Export local Feishu-like document fallback to simple PDF."},

    # Desktop: real GUI actions when pyautogui/PIL screen capture is available and explicitly enabled.
    "desktop.screenshot": {"risk": "A0", "implemented": True, "summary": "Capture desktop screenshot through PIL ImageGrab/pyautogui when host display is available."},
    "desktop.click": {"risk": "A3", "implemented": True, "summary": "Desktop click through pyautogui; requires OMNI_DESKTOP_ENABLE=1 or args.enable_desktop=true."},
    "desktop.type": {"risk": "A3", "implemented": True, "summary": "Desktop typing through pyautogui; requires OMNI_DESKTOP_ENABLE=1 or args.enable_desktop=true."},
    "desktop.hotkey": {"risk": "A3", "implemented": True, "summary": "Desktop hotkey through pyautogui; requires OMNI_DESKTOP_ENABLE=1 or args.enable_desktop=true."},
    "windows.desktop.screenshot": {"risk": "A0", "implemented": True, "summary": "Alias to desktop.screenshot."},
    "windows.desktop.click": {"risk": "A3", "implemented": True, "summary": "Alias to desktop.click."},
    "windows.desktop.type": {"risk": "A3", "implemented": True, "summary": "Alias to desktop.type."},
    "windows.desktop.hotkey": {"risk": "A3", "implemented": True, "summary": "Alias to desktop.hotkey."},
    "macos.desktop.screenshot": {"risk": "A0", "implemented": True, "summary": "Alias to desktop.screenshot."},
    "macos.desktop.click": {"risk": "A3", "implemented": True, "summary": "Alias to desktop.click."},
    "macos.desktop.type": {"risk": "A3", "implemented": True, "summary": "Alias to desktop.type."},
    "macos.desktop.hotkey": {"risk": "A3", "implemented": True, "summary": "Alias to desktop.hotkey."},
    "linux.desktop.screenshot": {"risk": "A0", "implemented": True, "summary": "Alias to desktop.screenshot."},
    "linux.desktop.click": {"risk": "A3", "implemented": True, "summary": "Alias to desktop.click."},
    "linux.desktop.type": {"risk": "A3", "implemented": True, "summary": "Alias to desktop.type."},
    "linux.desktop.hotkey": {"risk": "A3", "implemented": True, "summary": "Alias to desktop.hotkey."},

    # TTS: local OS TTS fallback where available.
    "audio.tts": {"risk": "A2", "implemented": True, "summary": "Create speech audio through local OS TTS engine if available; otherwise fail honestly."},
    "elevenlabs.tts.create": {"risk": "A2", "implemented": True, "summary": "Portable TTS fallback routed to local audio.tts. Real ElevenLabs needs API adapter."},
}
for _name, _patch in PORTABLE_APP_FALLBACKS.items():
    _merged = dict(ACTIONS.get(_name, {}))
    _merged.update(_patch)
    _merged.pop("adapter", None)
    ACTIONS[_name] = _merged

# ---------- v3.2 delivery kernel actions ----------
# These are the compact action names used by the PPT Skill through the single
# externally visible omni_body tool.  They must never disappear from validation
# merely because a dynamic submodule failed to import: that converts a precise
# capability fault into the misleading `unknown_action` symptom.
_REQUIRED_SINGLE_TOOL_ACTIONS: Dict[str, Dict[str, Any]] = {
    "skill.route": {"risk": "A0", "summary": "Route a task to a compact executable Skill."},
    "skill.get": {"risk": "A0", "summary": "Get compact Skill metadata."},
    "skill.read": {"risk": "A0", "summary": "Read the selected Skill procedure."},
    "template.apply": {"risk": "A2", "summary": "Apply a delivery template and emit its machine-readable design contract."},
    "qc.ppt.delivery_check": {"risk": "A0", "summary": "Run the PowerPoint delivery quality gate."},
    "repair.plan": {"risk": "A2", "summary": "Create a repair plan from quality-gate findings."},
    "deliverable.package": {"risk": "A2", "summary": "Package final deliverables and compact evidence."},
}


def _finalize_delivery_registry(
    actions: Dict[str, Dict[str, Any]],
    import_error: str = "",
) -> Tuple[Dict[str, Dict[str, Any]], Dict[str, str]]:
    """Keep required action names stable and expose partial-load diagnostics."""
    registry = dict(actions or {})
    errors: Dict[str, str] = {}
    for name, required in _REQUIRED_SINGLE_TOOL_ACTIONS.items():
        meta = registry.get(name)
        if isinstance(meta, dict):
            if not bool(meta.get("implemented", False)):
                errors[name] = str(meta.get("unavailable_reason") or "registered but not executable")
            continue
        reason = import_error or f"required action missing from delivery registry: {name}"
        registry[name] = {
            **required,
            "implemented": False,
            "unavailable_reason": reason,
            "capability_state": "unavailable",
        }
        errors[name] = reason
    return registry, errors


_DELIVERY_KERNEL_IMPORT_ERROR = ""
try:
    from .delivery_kernel import DELIVERY_ACTIONS as _LOADED_DELIVERY_ACTIONS, handle_delivery_action  # type: ignore
except Exception as _delivery_kernel_exc:  # keep portable core alive, but never hide degradation
    _DELIVERY_KERNEL_IMPORT_ERROR = f"{type(_delivery_kernel_exc).__name__}: {_delivery_kernel_exc}"
    _LOADED_DELIVERY_ACTIONS = {}
    handle_delivery_action = None  # type: ignore
DELIVERY_ACTIONS, _DELIVERY_CAPABILITY_ERRORS = _finalize_delivery_registry(
    _LOADED_DELIVERY_ACTIONS,
    _DELIVERY_KERNEL_IMPORT_ERROR,
)
ACTIONS.update(DELIVERY_ACTIONS)


def _required_capability_integrity() -> Dict[str, Any]:
    required = [*_REQUIRED_SINGLE_TOOL_ACTIONS, "pptx.create", "pptx.read"]
    status: Dict[str, Dict[str, Any]] = {}
    for name in required:
        meta = ACTIONS.get(name)
        status[name] = {
            "registered": isinstance(meta, dict),
            "implemented": bool(meta and meta.get("implemented", False)),
            "reason": str((meta or {}).get("unavailable_reason") or ""),
        }
    healthy = all(item["registered"] and item["implemented"] for item in status.values())
    return {
        "healthy": healthy,
        "required_action_count": len(status),
        "unavailable_actions": [name for name, item in status.items() if not item["implemented"]],
        "delivery_kernel_loaded": not bool(_DELIVERY_KERNEL_IMPORT_ERROR),
        "delivery_kernel_import_error": _DELIVERY_KERNEL_IMPORT_ERROR,
        "actions": status,
    }



@dataclass
class BodyRuntimeConfig:
    workspace: str = "."
    allow_absolute_paths: bool = False
    allow_shell: bool = False
    allow_python: bool = False
    require_confirmation_for_a4: bool = False
    max_text_read_chars: int = 200_000
    max_search_file_size_mb: int = 5
    audit_dir: str = ".omni_audit"
    backup_dir: str = ".omni_backups"
    trash_dir: str = ".omni_trash"
    emergency_audit_dir: str = ".tiangong_emergency_audit"
    ffmpeg_path: Optional[str] = None
    ffprobe_path: Optional[str] = None
    default_timeout_seconds: int = 0
    allowed_shell_commands: List[str] = field(default_factory=list)
    sandbox_enabled: bool = False
    sandbox_max_workspace_mb: int = 2048
    sandbox_max_changed_mb: int = 512
    sandbox_max_output_mb: int = 4
    sandbox_max_memory_mb: int = 2048
    sandbox_max_processes: int = 32
    fact_kernel_enabled: bool = True
    fact_ledger_root: str = ""
    run_id: str = ""
    request_id: str = ""
    generation: int = 0
    principal_scope_hash: str = ""
    skill_activation_sha256: str = ""
    gateway_url: str = ""
    session_id: str = ""
    step_id: str = ""
    task_node_id: str = ""
    host_verified_learning_intent: bool = False
    # D-21：网关随签发下发的用户指定路径根（用户原文提取）；空 = 仅工作区。
    user_path_roots: List[str] = field(default_factory=list)


class OmniBodyError(RuntimeError):
    pass


def _validate_cross_platform_filename(value: Any) -> str:
    """Validate one filename component identically on Linux and Windows."""
    if not isinstance(value, str) or not value:
        raise OmniBodyError("filename must be a non-empty string")
    if unicodedata.normalize("NFC", value) != value:
        raise OmniBodyError("filename must use NFC Unicode normalization")
    try:
        utf8_bytes = value.encode("utf-8", errors="strict")
        utf16_units = len(value.encode("utf-16-le", errors="strict")) // 2
    except UnicodeEncodeError as exc:
        raise OmniBodyError("filename contains an invalid Unicode scalar") from exc
    if len(utf8_bytes) > 255 or utf16_units > 255:
        raise OmniBodyError("filename exceeds the cross-platform component limit")
    if value != value.strip() or value.endswith((".", " ")):
        raise OmniBodyError("filename may not have surrounding spaces or a trailing dot")
    if value in {".", ".."}:
        raise OmniBodyError("relative path names are forbidden")
    forbidden_categories = {"Cc", "Cf", "Cs", "Zl", "Zp"}
    if any(unicodedata.category(char) in forbidden_categories for char in value):
        raise OmniBodyError("filename contains an invisible, directional, or control character")
    if any(char in value for char in '\\/:*?"<>|'):
        raise OmniBodyError("filename contains a path separator or reserved character")
    stem = value.split(".", 1)[0].upper()
    reserved = {"CON", "PRN", "AUX", "NUL"}
    reserved.update({f"COM{index}" for index in range(1, 10)})
    reserved.update({f"LPT{index}" for index in range(1, 10)})
    if stem in reserved:
        raise OmniBodyError("filename is reserved on Windows")
    return value


_PYTHON_EXECUTABLE_NAME = re.compile(r"^python(?:w)?(?:\d+(?:\.\d+)*)?(?:\.exe)?$", re.IGNORECASE)


def _is_python_interpreter(path: Path) -> bool:
    """Return True only for a real Python launcher, never the frozen backend exe."""
    try:
        resolved = path.expanduser().resolve(strict=True)
    except (OSError, RuntimeError):
        return False
    return resolved.is_file() and _PYTHON_EXECUTABLE_NAME.fullmatch(resolved.name) is not None


def _load_text_font(requested: Optional[str], size: int):
    """Load a scalable font, preferring Windows fonts with CJK coverage."""

    from PIL import ImageFont

    font_size = max(1, min(int(size), 4096))
    candidates: List[str] = []
    if requested:
        candidates.append(str(requested))
    if os.name == "nt":
        windows_root = str(
            os.environ.get("WINDIR") or os.environ.get("SystemRoot") or ""
        ).strip()
        if windows_root and Path(windows_root).is_absolute():
            fonts_dir = Path(windows_root) / "Fonts"
            candidates.extend(
                str(fonts_dir / name)
                for name in (
                    "msyh.ttc",
                    "msyhbd.ttc",
                    "simhei.ttf",
                    "simsun.ttc",
                    "Deng.ttf",
                    "arial.ttf",
                )
            )
    elif sys.platform == "darwin":
        fonts_dir = Path(os.path.sep) / "System" / "Library" / "Fonts"
        candidates.extend(
            str(fonts_dir / name)
            for name in (
                "PingFang.ttc",
                "Hiragino Sans GB.ttc",
            )
        )
    else:
        fonts_dir = Path(os.path.sep) / "usr" / "share" / "fonts"
        candidates.extend(
            str(fonts_dir / relative)
            for relative in (
                Path("opentype/noto/NotoSansCJK-Regular.ttc"),
                Path("truetype/wqy/wqy-zenhei.ttc"),
            )
        )
    candidates.extend(("DejaVuSans.ttf", "Arial.ttf"))
    seen: set[str] = set()
    for name in candidates:
        key = os.path.normcase(os.path.abspath(name))
        if key in seen:
            continue
        seen.add(key)
        try:
            return ImageFont.truetype(name, font_size)
        except (OSError, ValueError):
            continue
    return ImageFont.load_default()


def _run_windows_sapi_broker(
    powershell: str,
    *,
    text: str,
    output: Path,
    voice: str = "",
    rate: int = 0,
    timeout: int = 120,
) -> Dict[str, Any]:
    """Run fixed, injection-free Windows SAPI code outside the code sandbox.

    AppContainer intentionally cannot enumerate the host's installed speech
    voices. This broker is narrower than shell.run: every dynamic value is
    UTF-8/base64 data, the script is fixed, the output path was already
    workspace-confined by ``BodyRuntime._resolve``, and the environment is
    secret-free.
    """

    if os.name != "nt":
        raise OmniBodyError("Windows SAPI broker is only available on Windows")
    payloads = {
        "text": base64.b64encode(str(text).encode("utf-8")).decode("ascii"),
        "output": base64.b64encode(str(output).encode("utf-8")).decode("ascii"),
        "voice": base64.b64encode(str(voice or "").encode("utf-8")).decode("ascii"),
    }
    script = (
        "$ErrorActionPreference='Stop';"
        "$utf8=[Text.Encoding]::UTF8;"
        f"$text=$utf8.GetString([Convert]::FromBase64String('{payloads['text']}'));"
        f"$output=$utf8.GetString([Convert]::FromBase64String('{payloads['output']}'));"
        f"$requested=$utf8.GetString([Convert]::FromBase64String('{payloads['voice']}'));"
        "Add-Type -AssemblyName System.Speech;"
        "$speaker=New-Object System.Speech.Synthesis.SpeechSynthesizer;"
        "try{"
        "$installed=@($speaker.GetInstalledVoices() | Where-Object {$_.Enabled});"
        "if($requested){"
        "$chosen=$installed | Where-Object {"
        "$_.VoiceInfo.Name -eq $requested -or $_.VoiceInfo.Culture.Name -eq $requested"
        "} | Select-Object -First 1;"
        "if($chosen){$speaker.SelectVoice($chosen.VoiceInfo.Name)}"
        "} elseif($text -match '[\\u4e00-\\u9fff]'){"
        "$chosen=$installed | Where-Object {$_.VoiceInfo.Culture.Name -like 'zh*'} | Select-Object -First 1;"
        "if($chosen){$speaker.SelectVoice($chosen.VoiceInfo.Name)}"
        "};"
        f"$speaker.Rate={max(-10, min(int(rate), 10))};"
        "$speaker.SetOutputToWaveFile($output);"
        "$speaker.Speak($text)"
        "}finally{$speaker.Dispose()}"
    )
    encoded_script = base64.b64encode(script.encode("utf-16-le")).decode("ascii")
    command = [
        powershell,
        "-NoLogo",
        "-NoProfile",
        "-NonInteractive",
        "-ExecutionPolicy",
        "Bypass",
        "-EncodedCommand",
        encoded_script,
    ]
    started = time.monotonic()
    with tempfile.TemporaryDirectory(prefix="tiangong-sapi-") as temp_dir:
        env = subprocess_environment(sanitized_environment(os.environ, Path(temp_dir)))
        try:
            completed = subprocess.run(
                command,
                cwd=str(output.parent),
                env=env,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                shell=False,
                timeout=max(1, min(int(timeout), 300)),
                creationflags=subprocess.CREATE_NO_WINDOW,
                check=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise OmniBodyError("Windows SAPI broker timed out") from exc
    stdout = decode_portable_bytes(
        completed.stdout,
        source="Windows SAPI stdout",
        allow_legacy_windows=True,
    ).text
    stderr = decode_portable_bytes(
        completed.stderr,
        source="Windows SAPI stderr",
        allow_legacy_windows=True,
    ).text
    return {
        "returncode": int(completed.returncode),
        "stdout": _bounded_subprocess_text(stdout),
        "stderr": _bounded_subprocess_text(stderr),
        "ok": completed.returncode == 0,
        "containment": "trusted-windows-sapi-broker",
        "network": "not-used",
        "elapsed_seconds": round(time.monotonic() - started, 3),
    }


def _resolve_python_interpreter() -> str:
    """Resolve the interpreter used by python.run in source and frozen builds.

    In a PyInstaller build ``sys.executable`` is tiangong-backend.exe.  Reusing
    it as a Python interpreter starts a second backend process and can steal the
    live port, so it must never be accepted merely because it exists.
    """
    candidates: List[Path] = []
    configured = str(os.environ.get("TIANGONG_PYTHON_EXECUTABLE") or "").strip()
    if configured:
        candidates.append(Path(configured))

    executable = Path(sys.executable)
    candidates.append(executable)

    # Packaged layout: app/backend/tiangong-backend/tiangong-backend.exe and
    # app/life-service/runtime314/python.exe share the same application root.
    for parent in executable.expanduser().absolute().parents:
        if parent.name.casefold() == "app":
            candidates.append(parent / "life-service" / "runtime314" / "python.exe")
            break
    for parent in Path(__file__).resolve().parents:
        if parent.name.casefold() == "app":
            candidates.append(parent / "life-service" / "runtime314" / "python.exe")
            break

    for command in ("python", "python3"):
        discovered = shutil.which(command)
        if discovered:
            candidates.append(Path(discovered))

    seen: set[str] = set()
    for candidate in candidates:
        key = os.path.normcase(os.path.abspath(str(candidate)))
        if key in seen:
            continue
        seen.add(key)
        if _is_python_interpreter(candidate):
            return str(candidate.expanduser().resolve(strict=True))
    raise OmniBodyError(
        "No trusted Python interpreter is available; configure TIANGONG_PYTHON_EXECUTABLE."
    )


def _bounded_subprocess_text(value: Any, limit: int | None = None) -> str:
    """Return complete subprocess output even when a stream is ``None``."""
    if value is None:
        return ""
    if isinstance(value, bytes):
        value = decode_portable_bytes(value, source="subprocess output", allow_legacy_windows=True).text
    elif not isinstance(value, str):
        value = str(value)
    if limit is None:
        return value
    return value[-max(0, int(limit)) :]


def _capability_prefix_counts(action_names: Iterable[str]) -> Dict[str, int]:
    """Summarize a large action registry without injecting every name."""
    counts: Dict[str, int] = {}
    for name in action_names:
        prefix = str(name).split(".", 1)[0] or "other"
        counts[prefix] = counts.get(prefix, 0) + 1
    return dict(sorted(counts.items()))


class BodyRuntime:
    def __init__(self, config: Optional[BodyRuntimeConfig] = None):
        self.config = config or BodyRuntimeConfig()
        self.workspace = Path(self.config.workspace).expanduser().resolve()
        self.workspace.mkdir(parents=True, exist_ok=True)
        self._workspace_lock = _workspace_lock_for(self.workspace)
        state_root_raw = str(os.environ.get("TIANGONG_OMNI_BODY_STATE_ROOT") or "").strip()
        state_root = Path(state_root_raw).expanduser().resolve() if state_root_raw else None
        audit_dir = str(state_root / "audit") if state_root and self.config.audit_dir == ".omni_audit" else self.config.audit_dir
        backup_dir = str(state_root / "backups") if state_root and self.config.backup_dir == ".omni_backups" else self.config.backup_dir
        trash_dir = str(state_root / "trash") if state_root and self.config.trash_dir == ".omni_trash" else self.config.trash_dir
        self.audit_dir = self._resolve_special(audit_dir)
        self.emergency_audit_dir = self._resolve_special(self.config.emergency_audit_dir)
        self.backup_dir = self._resolve_special(backup_dir)
        self.trash_dir = self._resolve_special(trash_dir)
        sandbox_root = (state_root / "sandboxes") if state_root else (self.workspace / ".tiangong_sandboxes")
        self.sandbox = SandboxRunner(
            self.workspace, sandbox_root, self.trash_dir,
            SandboxLimits(
                timeout_seconds=max(1, int(self.config.default_timeout_seconds)),
                max_workspace_bytes=max(1, int(self.config.sandbox_max_workspace_mb)) * 1024 * 1024,
                max_changed_bytes=max(1, int(self.config.sandbox_max_changed_mb)) * 1024 * 1024,
                max_output_bytes=max(1, int(self.config.sandbox_max_output_mb)) * 1024 * 1024,
                max_memory_bytes=max(128, int(self.config.sandbox_max_memory_mb)) * 1024 * 1024,
                max_processes=max(1, int(self.config.sandbox_max_processes)),
            ),
        )
        self._state_init_warnings: List[Dict[str, str]] = []
        self._execution_state = threading.local()
        self._audit_target = self.audit_dir
        for name, directory in (("audit", self.audit_dir), ("emergency_audit", self.emergency_audit_dir), ("backup", self.backup_dir), ("trash", self.trash_dir)):
            try:
                directory.mkdir(parents=True, exist_ok=True)
            except Exception as exc:
                self._state_init_warnings.append({"state": name, "path": str(directory), "error": f"{type(exc).__name__}: {exc}"})
        self.ffmpeg = self.config.ffmpeg_path or shutil.which("ffmpeg")
        if not self.ffmpeg:
            try:
                import imageio_ffmpeg  # type: ignore

                portable_ffmpeg = str(imageio_ffmpeg.get_ffmpeg_exe() or "").strip()
                if portable_ffmpeg and Path(portable_ffmpeg).is_file():
                    self.ffmpeg = portable_ffmpeg
            except Exception:
                # ``system.health`` reports the missing adapter explicitly;
                # optional dependency discovery must never prevent startup.
                self.ffmpeg = None
        self.ffprobe = self.config.ffprobe_path or shutil.which("ffprobe")
        self.capability_manifest = compile_manifest(ACTIONS, self.__class__, dynamic_actions=set(DELIVERY_ACTIONS))
        self.capability_integrity = _required_capability_integrity()
        run_id = str(self.config.run_id or self.config.request_id or f"body_{os.getpid()}_{uuid.uuid4().hex[:16]}")
        ledger_root = str(self.config.fact_ledger_root or (state_root / "fact_ledger" if state_root else self.workspace / ".tiangong" / "fact_ledger"))
        self.fact_kernel = FactExecutionKernel(
            self.workspace, ledger_root, run_id,
            request_id=str(self.config.request_id or run_id),
            session_id=str(self.config.session_id or ""),
        )

    def run(self, action: str, target: Optional[str] = None, args: Optional[Dict[str, Any]] = None, **kwargs: Any) -> Dict[str, Any]:
        payload = dict(args or {})
        payload.update(kwargs)
        validation = validate_tool_request(
            action,
            target,
            payload,
            workspace=self.workspace,
            available_actions=ACTIONS,
            user_roots=getattr(self.config, "user_path_roots", None) or [],
        )
        if not validation.get("ok"):
            return {
                **validation,
                "success": False,
                "risk_level": ACTIONS.get(str(validation.get("action") or ""), {}).get("risk"),
            }
        action = str(validation.get("action") or "")
        target = str(validation.get("target") or "")
        payload = dict(validation.get("args") or {})
        if self.config.fact_kernel_enabled and not fact_execution_active():
            expected_version = str(payload.pop("expected_version", "") or "")
            idempotency_key = str(payload.pop("idempotency_key", "") or "")
            step_id = str(payload.pop("__step_id", "") or self.config.step_id or "")
            task_node_id = str(payload.pop("__task_node_id", "") or self.config.task_node_id or step_id)
            result = self.fact_kernel.execute(
                action, target, payload,
                lambda op_id: self._run_legacy(action, target, payload, _op_id=op_id),
                step_id=step_id, task_node_id=task_node_id,
                expected_version=expected_version, idempotency_key=idempotency_key,
            )
            transaction = dict(result.get("fact_transaction") or {})
            recreatable_outputs = {"docx.create", "word.create", "pptx.create", "mindmap.create"}
            if (
                action in recreatable_outputs
                and transaction.get("idempotent_replay") is True
                and not self._recreatable_output_valid(action, target, payload, result)
            ):
                # An OBSERVED journal record proves that a previous invocation
                # completed, not that its external artifact still exists.  A
                # successful replay with a missing target is therefore stale.
                # Re-run the deterministic generator as a new recovery effect
                # and record that recovery instead of returning phantom success.
                recovery_key = f"{idempotency_key or transaction.get('operation_id') or action}:missing-output-recovery:{uuid.uuid4().hex}"
                recovered = self.fact_kernel.execute(
                    action, target, payload,
                    lambda op_id: self._run_legacy(action, target, payload, _op_id=op_id),
                    step_id=step_id, task_node_id=task_node_id,
                    expected_version=expected_version, idempotency_key=recovery_key,
                )
                recovered["stale_replay_recovered"] = True
                recovered["stale_replay_operation_id"] = str(transaction.get("operation_id") or "")
                if not self._recreatable_output_valid(action, target, payload, recovered):
                    recovered.update({
                        "success": False,
                        "ok": False,
                        "status": "RECOVERY_OUTPUT_INVALID",
                        "error": "recreated output failed structural or digest validation",
                    })
                return recovered
            return result
        return self._run_legacy(action, target, payload)

    def _run_legacy(self, action: str, target: Optional[str] = None, args: Optional[Dict[str, Any]] = None, *, _op_id: Optional[str] = None) -> Dict[str, Any]:
        meta = ACTIONS.get(action, {})
        if meta.get("risk") in {"A1", "A2", "A3", "A4"}:
            with _workspace_mutation_guard(self.workspace, self._workspace_lock):
                return self._run_unlocked(action, target, args, _op_id=_op_id)
        return self._run_unlocked(action, target, args, _op_id=_op_id)

    def _run_unlocked(self, action: str, target: Optional[str] = None, args: Optional[Dict[str, Any]] = None, *, _op_id: Optional[str] = None, **kwargs: Any) -> Dict[str, Any]:
        """Single entry point implementation.

        Args:
            action: explicit action name, e.g. "file.read", "video.cut".
            target: primary file/directory/resource path, relative to workspace by default.
            args: action-specific arguments.
            **kwargs: merged into args for convenience.
        """
        args = dict(args or {})
        args.update(kwargs)
        started = time.time()
        op_id = str(_op_id or self._new_op_id(action))
        try:
            if action not in ACTIONS:
                raise OmniBodyError(f"Unknown action: {action}")
            meta = ACTIONS[action]
            if meta.get("risk") == "A5":
                return self._adapter_or_blocked(action, target, args, reason="A5 hard-gate: blocked by default")
            if not meta.get("implemented", False):
                return self._adapter_or_blocked(
                    action,
                    target,
                    args,
                    reason=str(meta.get("unavailable_reason") or "portable core does not implement this adapter"),
                )
            if meta.get("risk") == "A4":
                audit_error = self._audit_preflight_error()
                if audit_error:
                    return {
                        "success": False,
                        "op_id": op_id,
                        "action": action,
                        "target": target,
                        "risk_level": meta.get("risk"),
                        "error_type": "AuditUnavailable",
                        "message": audit_error,
                        "audit_required": True,
                        "audit_persisted": False,
                        "elapsed_seconds": round(time.time() - started, 3),
                    }

            alias_to = meta.get("alias_to")
            if alias_to:
                routed_args = dict(meta.get("default_args") or {})
                routed_args.update(args)
                routed = self.run(str(alias_to), target, routed_args)
                if isinstance(routed, dict):
                    routed = dict(routed)
                    routed["routed_from"] = action
                    routed["routed_to"] = str(alias_to)
                    routed["app_id"] = meta.get("app_id")
                    routed["tool_group"] = meta.get("tool_group")
                    routed["action"] = action
                    routed.setdefault("success", bool(routed.get("success", False)))
                    return routed
                return {"success": False, "action": action, "routed_to": str(alias_to), "message": "Alias route returned non-dict result."}

            previous_action = getattr(self._execution_state, "action", None)
            previous_risk = getattr(self._execution_state, "risk", None)
            self._execution_state.action = action
            self._execution_state.risk = str(meta.get("risk") or "A0")
            try:
                if action in globals().get("DELIVERY_ACTIONS", {}):
                    if globals().get("handle_delivery_action") is None:
                        raise OmniBodyError(f"Delivery kernel unavailable for {action}")
                    result = globals()["handle_delivery_action"](self, op_id, action, target, args)
                else:
                    handler_name = "_action_" + action.replace(".", "_")
                    handler = getattr(self, handler_name, None)
                    if handler is None:
                        raise OmniBodyError(f"No handler implemented for {action}")
                    result = handler(op_id, target, args)
            finally:
                self._execution_state.action = previous_action
                self._execution_state.risk = previous_risk
            if not isinstance(result, dict):
                result = {"result": result}
            if result.get("ok") is False:
                result["success"] = False
            else:
                result.setdefault("success", True)
            result.setdefault("op_id", op_id)
            result.setdefault("action", action)
            result.setdefault("risk_level", meta.get("risk"))
            result.setdefault("elapsed_seconds", round(time.time() - started, 3))
            self._safe_write_audit(op_id, action, target, args, result)
            return result
        except Exception as exc:  # explicit structured error, not swallowed
            result = {
                "success": False,
                "op_id": op_id,
                "action": action,
                "target": target,
                "risk_level": ACTIONS.get(action, {}).get("risk"),
                "error_type": exc.__class__.__name__,
                "message": str(exc),
                "traceback": traceback.format_exc(limit=5),
                "elapsed_seconds": round(time.time() - started, 3),
            }
            self._safe_write_audit(op_id, action, target, args, result)
            return result

    # ---------- path, audit, snapshot, rollback ----------

    def _resolve_special(self, path: str) -> Path:
        p = Path(path).expanduser()
        if not p.is_absolute():
            p = self.workspace / p
        return p.resolve()

    def _is_inside(self, path: Path, root: Path) -> bool:
        try:
            path.resolve().relative_to(root.resolve())
            return True
        except (OSError, RuntimeError, ValueError):
            return False

    def _resolve(self, path: Optional[str], must_exist: bool = False) -> Path:
        if path is None or not str(path).strip():
            raise OmniBodyError("path target is required; workspace root inference is forbidden")
        raw = str(path)
        if unicodedata.normalize("NFC", raw) != raw or "\x00" in raw:
            raise OmniBodyError("path must be NFC and contain no NUL")
        folded = raw.replace("/", "\\")
        if folded.startswith(("\\\\.\\", "\\\\?\\")):
            raise OmniBodyError("Windows device paths are forbidden")
        candidate = Path(raw).expanduser()
        user_roots = getattr(self.config, "user_path_roots", None) or []
        _user_allowed = globals().get("_user_path_root_allowed")
        if _user_allowed is None:  # AST 抽取/单文件加载环境：模块级不可见时的内联兜底（同一口径）
            def _user_allowed(resolved_text: str, roots) -> bool:
                if not roots:
                    return False
                resolved = Path(str(resolved_text))
                text = os.path.normcase(str(resolved))
                for prefix in (
                    os.path.normcase(os.environ.get("SystemRoot") or r"C:\Windows"),
                    os.path.normcase(os.environ.get("ProgramFiles") or r"C:\Program Files"),
                    os.path.normcase(os.environ.get("ProgramFiles(x86)") or r"C:\Program Files (x86)"),
                    os.path.normcase(os.environ.get("ProgramData") or r"C:\ProgramData"),
                ):
                    if text == prefix or text.startswith(prefix + os.sep):
                        return False
                if len(text) <= 3:
                    return False
                if any(part.casefold() in {".ssh", ".aws", ".gnupg", ".azure", ".config"} for part in resolved.parts):
                    return False
                if resolved.name.casefold().endswith(".env"):
                    return False
                for root in roots:
                    try:
                        resolved.relative_to(Path(str(root)).resolve(strict=False))
                        return True
                    except (ValueError, OSError):
                        continue
                return False
        if not self.config.allow_absolute_paths:
            if candidate.is_absolute():
                # An absolute spelling of a path already inside the selected
                # workspace grants no extra authority. Normalize lexically,
                # then retain the reparse and final resolved-containment gates
                # below so junctions/symlinks cannot turn this into an escape.
                candidate = Path(os.path.abspath(str(candidate)))
                normalized_root = os.path.normcase(
                    unicodedata.normalize("NFC", str(self.workspace.resolve()))
                )
                normalized_candidate = os.path.normcase(
                    unicodedata.normalize("NFC", str(candidate))
                )
                try:
                    absolute_inside_workspace = (
                        os.path.commonpath((normalized_root, normalized_candidate))
                        == normalized_root
                    )
                except (OSError, ValueError):
                    absolute_inside_workspace = False
                # D-21：用户本轮明确指定的路径根（网关随签发下发）视为已授权，
                # 与工作区内同权；硬禁区（系统目录/凭据目录/.env）不豁免。
                if not absolute_inside_workspace and not _user_allowed(
                    str(candidate), user_roots
                ):
                    raise OmniBodyError("absolute paths require an exact signed A4 capability")
            elif (
                raw.startswith(("/", "\\\\"))
                or re.match(r"^[A-Za-z]:[\\/]", raw)
                or ".." in candidate.parts
            ):
                raise OmniBodyError("path traversal, UNC, and rooted paths are forbidden")
        if not candidate.is_absolute():
            candidate = self.workspace / candidate
        # Reject every existing reparse component before resolve() can erase
        # evidence of a symlink or junction hop.
        cursor = candidate
        while True:
            if cursor.exists() or cursor.is_symlink():
                is_junction = bool(getattr(cursor, "is_junction", lambda: False)())
                if cursor.is_symlink() or is_junction:
                    raise OmniBodyError(f"reparse path component is forbidden: {cursor}")
                try:
                    attributes = int(getattr(cursor.stat(follow_symlinks=False), "st_file_attributes", 0))
                except (OSError, TypeError):
                    attributes = 0
                if attributes & 0x400:
                    raise OmniBodyError(f"Windows reparse point is forbidden: {cursor}")
            if cursor == cursor.parent or self._is_inside(cursor, self.workspace) and cursor == self.workspace:
                break
            cursor = cursor.parent
        p = candidate.resolve(strict=False)
        normalized_root = os.path.normcase(unicodedata.normalize("NFC", str(self.workspace.resolve())))
        normalized_path = os.path.normcase(unicodedata.normalize("NFC", str(p)))
        try:
            inside = os.path.commonpath((normalized_root, normalized_path)) == normalized_root
        except ValueError:
            inside = False
        if not self.config.allow_absolute_paths and not inside and not _user_allowed(str(p), user_roots):
            raise OmniBodyError(f"Path escapes workspace: {path}")
        if must_exist and not p.exists():
            raise FileNotFoundError(str(p))
        execution_state = getattr(self, "_execution_state", None)
        # A directly constructed boundary has no action context.  Treat that
        # compatibility/security-test path as mutating by default; only the
        # real runtime may explicitly mark the current action read-only.
        current_risk = str(getattr(execution_state, "risk", "A3") or "A3")
        if p.exists() and p.is_file() and p.stat().st_nlink > 1 and current_risk in {"A1", "A2", "A3", "A4"}:
            raise OmniBodyError("hard-linked files cannot be mutated without object-grant materialization")
        return p

    def _rel(self, p: Path) -> str:
        try:
            return p.resolve().relative_to(self.workspace).as_posix()
        except ValueError:
            # File outside workspace — strip drive letter and use safe path
            raw = str(p.resolve())
            # Remove Windows drive letter prefix (e.g. "C:\" → "")
            if len(raw) >= 2 and raw[1] == ":":
                raw = raw[2:]
            # Replace backslashes and colons that can't appear in file paths
            safe = raw.replace("\\", "/").lstrip("/")
            return safe

    def _new_op_id(self, action: str) -> str:
        safe_action = re.sub(r"[^a-zA-Z0-9_.-]+", "_", action)[:48]
        # Millisecond-only ids collide under multi-agent/concurrent execution.
        # time_ns preserves ordering context; UUID guarantees uniqueness across
        # threads and processes that share the same workspace.
        return f"{time.time_ns()}_{os.getpid()}_{uuid.uuid4().hex[:12]}_{safe_action}"

    def _write_audit(self, op_id: str, action: str, target: Optional[str], args: Dict[str, Any], result: Dict[str, Any]) -> None:
        with _workspace_mutation_guard(self.workspace, self._workspace_lock):
            audit_target = self._select_audit_target()
            audit_target.mkdir(parents=True, exist_ok=True)
            record = {
                "op_id": op_id,
                "ts": time.time(),
                "time_local": time.strftime("%Y-%m-%d %H:%M:%S"),
                "workspace": str(self.workspace),
                "action": action,
                "target": target,
                "args_safe": self._redact_args(args),
                "result": self._shorten_for_log(result),
                "audit_fallback": audit_target == self.emergency_audit_dir,
            }
            json_path = audit_target / f"{op_id}.json"
            json_path.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")
            with (audit_target / "audit.jsonl").open("a", encoding="utf-8", newline="\n") as f:
                f.write(json.dumps(record, ensure_ascii=False) + "\n")

    @staticmethod
    def _probe_writable_directory(directory: Path) -> str:
        probe_path = directory / f".audit-probe-{os.getpid()}-{time.time_ns()}"
        try:
            directory.mkdir(parents=True, exist_ok=True)
            with probe_path.open("x", encoding="utf-8", newline="\n") as probe:
                probe.write("ok")
                probe.flush()
                os.fsync(probe.fileno())
            probe_path.unlink()
            return ""
        except Exception as exc:
            return f"{directory}: {type(exc).__name__}: {exc}"
        finally:
            try:
                probe_path.unlink(missing_ok=True)
            except OSError:
                pass

    def _select_audit_target(self) -> Path:
        primary_error = self._probe_writable_directory(self.audit_dir)
        if not primary_error:
            self._audit_target = self.audit_dir
            return self.audit_dir
        fallback_error = self._probe_writable_directory(self.emergency_audit_dir)
        if not fallback_error:
            self._audit_target = self.emergency_audit_dir
            return self.emergency_audit_dir
        raise OmniBodyError(
            "audit directories are not writable: primary=" + primary_error + "; fallback=" + fallback_error
        )

    def _audit_preflight_error(self) -> str:
        try:
            self._select_audit_target()
            return ""
        except Exception as exc:
            return str(exc)

    def _safe_write_audit(self, op_id: str, action: str, target: Optional[str], args: Dict[str, Any], result: Dict[str, Any]) -> None:
        try:
            self._write_audit(op_id, action, target, args, result)
            result["audit_persisted"] = True
        except Exception as exc:
            warning = {"error_type": type(exc).__name__, "message": str(exc), "path": str(self.audit_dir)}
            result["audit_persisted"] = False
            result["audit_warning"] = warning
            result.setdefault("warnings", []).append({"kind": "audit_not_persisted", **warning})

    def _redact_args(self, args: Dict[str, Any]) -> Dict[str, Any]:
        redacted = {}
        for k, v in args.items():
            if any(tok in k.lower() for tok in ["password", "token", "secret", "key", "cookie"]):
                redacted[k] = "<redacted>"
            elif isinstance(v, str) and len(v) > 2000:
                redacted[k] = v[:2000] + f"...<truncated {len(v)-2000} chars>"
            else:
                redacted[k] = v
        return redacted

    def _shorten_for_log(self, obj: Any, max_chars: int = 4000) -> Any:
        s = json.dumps(obj, ensure_ascii=False, default=str)
        if len(s) <= max_chars:
            return obj
        return {"truncated": True, "preview": s[:max_chars]}

    def _snapshot(self, op_id: str, paths: Iterable[Path]) -> List[Dict[str, Any]]:
        snapshots: List[Dict[str, Any]] = []
        seen = set()
        for raw in paths:
            p = raw.resolve()
            if str(p) in seen:
                continue
            seen.add(str(p))
            snap: Dict[str, Any] = {
                "path": str(p),
                "rel_path": self._rel(p),
                "existed": p.exists(),
                "kind": "dir" if p.is_dir() else "file" if p.is_file() else "other",
                "backup_path": None,
            }
            if p.exists():
                backup = self.backup_dir / op_id / self._rel(p)
                backup.parent.mkdir(parents=True, exist_ok=True)
                if p.is_dir():
                    if backup.exists():
                        shutil.rmtree(backup)
                    shutil.copytree(p, backup, ignore=shutil.ignore_patterns(self.config.audit_dir, self.config.backup_dir, self.config.trash_dir))
                elif p.is_file():
                    try:
                        shutil.copy2(p, backup)
                    except (PermissionError, OSError) as e:
                        # File locked by another process — fall back to content copy
                        try:
                            with open(p, "rb") as fsrc:
                                with open(backup, "wb") as fdst:
                                    while True:
                                        chunk = fsrc.read(1024 * 1024)
                                        if not chunk:
                                            break
                                        fdst.write(chunk)
                            snap["backup_fallback"] = True
                        except (PermissionError, OSError):
                            # Even content read is blocked — skip this file, don't crash
                            snap["backup_skipped"] = True
                            snap["backup_error"] = str(e)
                            snapshots.append(snap)
                            continue
                else:
                    # Ignore unusual sockets/devices; workspace normally should not contain them.
                    snap["existed"] = False
                snap["backup_path"] = str(backup) if backup.exists() else None
            snapshots.append(snap)
        (self.backup_dir / op_id).mkdir(parents=True, exist_ok=True)
        (self.backup_dir / op_id / "snapshots.json").write_text(json.dumps(snapshots, ensure_ascii=False, indent=2), encoding="utf-8")
        return snapshots

    def _remove_path(self, p: Path) -> None:
        if not p.exists():
            return
        if p.is_dir():
            shutil.rmtree(p)
        else:
            p.unlink()

    def _restore_snapshots(self, snapshots: List[Dict[str, Any]]) -> Dict[str, Any]:
        restored = []
        # restore most nested paths first to reduce parent/child conflicts
        for snap in sorted(snapshots, key=lambda s: len(s.get("path", "")), reverse=True):
            p = Path(snap["path"])
            if (
                not self.config.allow_absolute_paths
                and not self._is_inside(p, self.workspace)
                and not _user_path_root_allowed(str(p), getattr(self.config, "user_path_roots", None) or [])
            ):
                raise OmniBodyError(f"Rollback path escapes workspace: {p}")
            if snap.get("existed") and snap.get("backup_path"):
                backup = Path(snap["backup_path"])
                self._remove_path(p)
                p.parent.mkdir(parents=True, exist_ok=True)
                if backup.is_dir():
                    shutil.copytree(backup, p)
                else:
                    shutil.copy2(backup, p)
                restored.append({"path": str(p), "mode": "restored_backup"})
            else:
                self._remove_path(p)
                restored.append({"path": str(p), "mode": "removed_new_path"})
        return {"restored": restored}

    def _sha256(self, p: Path) -> str:
        h = hashlib.sha256()
        with p.open("rb") as f:
            for chunk in iter(lambda: f.read(1024 * 1024), b""):
                h.update(chunk)
        return h.hexdigest()

    def _file_evidence(self, p: Path, *, with_hash: bool = True) -> Dict[str, Any]:
        exists = p.exists()
        result = {
            "path": str(p),
            "rel_path": self._rel(p),
            "exists": exists,
            "is_file": p.is_file() if exists else False,
            "is_dir": p.is_dir() if exists else False,
            "size_bytes": p.stat().st_size if exists and p.is_file() else None,
        }
        if with_hash and exists and p.is_file() and p.stat().st_size <= 200 * 1024 * 1024:
            result["sha256"] = self._sha256(p)
        return result

    def _recreatable_output_valid(self, action: str, target: Optional[str], args: Dict[str, Any], result: Dict[str, Any]) -> bool:
        try:
            output = self._resolve(target or args.get("output"))
            if not output.is_file() or output.is_symlink() or output.stat().st_size <= 0:
                return False

            def evidence_matches(path: Path, evidence: Any) -> bool:
                if not isinstance(evidence, dict):
                    return True
                expected_size = evidence.get("size_bytes", evidence.get("bytes"))
                if isinstance(expected_size, int) and not isinstance(expected_size, bool) and path.stat().st_size != expected_size:
                    return False
                expected_hash = str(evidence.get("sha256") or "").lower()
                if re.fullmatch(r"[0-9a-f]{64}", expected_hash):
                    digest = hashlib.sha256()
                    with path.open("rb") as handle:
                        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                            digest.update(chunk)
                    if digest.hexdigest() != expected_hash:
                        return False
                return True

            if not evidence_matches(output, result.get("output")):
                return False
            if action in {"docx.create", "word.create", "pptx.create"}:
                with zipfile.ZipFile(output, "r") as archive:
                    if archive.testzip() is not None:
                        return False
                    names = set(archive.namelist())
                if "[Content_Types].xml" not in names:
                    return False
                if action in {"docx.create", "word.create"}:
                    return "word/document.xml" in names
                return "ppt/presentation.xml" in names and any(
                    name.startswith("ppt/slides/slide") and name.endswith(".xml") for name in names
                )
            if action == "mindmap.create":
                markdown = output.read_text(encoding="utf-8", errors="strict")
                if "```mermaid" not in markdown or "\nmindmap\n" not in markdown or markdown.count("root((") != 1:
                    return False
                if args.get("opml"):
                    import xml.etree.ElementTree as ET
                    opml_path = output.with_suffix(".opml")
                    if not opml_path.is_file() or opml_path.is_symlink() or not evidence_matches(opml_path, result.get("opml")):
                        return False
                    opml_root = ET.parse(opml_path).getroot()
                    if opml_root.tag != "opml" or opml_root.find("body") is None:
                        return False
                return True
            return False
        except (OSError, UnicodeError, ValueError, zipfile.BadZipFile):
            return False

    def _require_ffmpeg(self) -> str:
        if not self.ffmpeg:
            raise OmniBodyError("ffmpeg not found. Install ffmpeg or configure ffmpeg_path.")
        return self.ffmpeg

    def _require_ffprobe(self) -> str:
        if not self.ffprobe:
            raise OmniBodyError("ffprobe not found. Install ffprobe or configure ffprobe_path.")
        return self.ffprobe

    def _normalize_shell_command(self, command: Any, *, action: str) -> Tuple[List[str] | str, str]:
        if isinstance(command, str):
            parsed = shlex.split(command, posix=os.name != "nt")
            if not parsed:
                raise OmniBodyError(f"{action} command is empty")
            executable = Path(parsed[0]).name
            if os.name == "nt":
                # Keep the command body plain until SandboxRunner rewrites
                # absolute workspace paths. It then encodes the command into a
                # UTF-8-initialized PowerShell launcher, preserving cmd syntax
                # for quotes and metacharacters without an outer cmd layer.
                command_line = [WINDOWS_UTF8_SHELL_MARKER, command]
            else:
                command_line = ["/bin/sh", "-lc", command]
            return command_line, executable

        command_line = list(command)
        if not command_line:
            raise OmniBodyError(f"{action} command is empty")
        return command_line, Path(command_line[0]).name

    def _run_subprocess(
        self,
        cmd: List[str] | str,
        timeout: Optional[int] = None,
        shell: bool = False,
        cwd: Optional[Path] = None,
        op_id: str = "",
    ) -> Dict[str, Any]:
        if shell:
            raise OmniBodyError("shell=True is forbidden; commands must be explicit argument vectors")
        timeout = self.config.default_timeout_seconds if timeout is None else timeout
        run_cwd = Path(cwd) if cwd is not None else self.workspace
        if not self.config.sandbox_enabled:
            before_files = (
                _sandbox_workspace_snapshot(run_cwd)
                if run_cwd.is_dir()
                else {}
            )
            prepared = (
                _prepare_windows_utf8_shell_command(cmd, cwd=run_cwd)
                if os.name == "nt"
                else cmd
            )
            started = time.monotonic()
            completed = subprocess.run(
                prepared,
                cwd=str(run_cwd),
                env=subprocess_environment(os.environ),
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                shell=False,
                timeout=None,
                check=False,
            )
            stdout = decode_portable_bytes(
                completed.stdout or b"",
                source="host subprocess stdout",
                allow_legacy_windows=True,
            )
            stderr = decode_portable_bytes(
                completed.stderr or b"",
                source="host subprocess stderr",
                allow_legacy_windows=True,
            )
            after_files = (
                _sandbox_workspace_snapshot(run_cwd)
                if run_cwd.is_dir()
                else {}
            )
            changed_files = sorted(
                path for path, evidence in after_files.items()
                if before_files.get(path) != evidence
            )
            deleted_files = sorted(set(before_files).difference(after_files))
            return {
                "returncode": int(completed.returncode),
                "stdout": stdout.text,
                "stderr": stderr.text,
                "stdout_encoding": stdout.encoding,
                "stderr_encoding": stderr.encoding,
                "legacy_output_encoding": bool(
                    stdout.legacy_fallback or stderr.legacy_fallback
                ),
                "ok": int(completed.returncode) == 0,
                "containment": "gateway_a5_host_execution",
                "elapsed_seconds": round(time.monotonic() - started, 3),
                "timeout_disabled": True,
                "changed_files": changed_files,
                "deleted_files": deleted_files,
                "changed_bytes": sum(after_files[path][0] for path in changed_files),
            }
        result = self.sandbox.run(cmd, cwd=run_cwd, timeout_seconds=timeout, op_id=op_id)
        result["stdout"] = _bounded_subprocess_text(result.get("stdout"))
        result["stderr"] = _bounded_subprocess_text(result.get("stderr"))
        return result

    def _adapter_or_blocked(self, action: str, target: Optional[str], args: Dict[str, Any], reason: str) -> Dict[str, Any]:
        meta = ACTIONS.get(action, {})
        result = {
            "success": False,
            "action": action,
            "target": target,
            "risk_level": meta.get("risk"),
            "ok": False,
            "outcome": "NOT_EXECUTED",
            "implemented": False,
            "requires_adapter": meta.get("adapter"),
            "reason": reason,
            "message": meta.get("summary"),
            "evidence": {},
        }
        if meta.get("unavailable_reason"):
            result.update({
                "error_type": "CapabilityUnavailable",
                "capability_state": "unavailable",
                "retryable": False,
            })
        return result

    # ---------- system ----------

    def _action_system_capabilities(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        include_actions = bool(args.get("include_actions", False))
        exposed_only = bool(args.get("executable_only", True))
        manifest = self.capability_manifest.to_dict(exposed_only=exposed_only)
        definitions = dict(manifest.get("capabilities") or {})
        action_names = sorted(definitions)
        dynamic_action_names = list(manifest.get("dynamic_actions") or [])
        compact_manifest = {
            "runtime_class": manifest.get("runtime_class"),
            "action_count": len(action_names),
            "dynamic_action_count": len(dynamic_action_names),
        }
        if include_actions:
            compact_manifest["dynamic_actions"] = dynamic_action_names
            actions = {
                name: {
                    "risk": row.get("risk"),
                    "implemented": bool(row.get("implemented")),
                    "effect": row.get("effect"),
                    "summary": row.get("summary"),
                }
                for name, row in definitions.items()
            }
        else:
            actions = {
                "count": len(action_names),
                "names_omitted": True,
                "prefix_counts": _capability_prefix_counts(action_names),
            }
        return {
            "capability_manifest": compact_manifest,
            "actions": actions,
            "required_capability_integrity": {
                "healthy": bool(self.capability_integrity.get("healthy")),
                "unavailable_actions": list(self.capability_integrity.get("unavailable_actions") or []),
                "delivery_kernel_loaded": bool(self.capability_integrity.get("delivery_kernel_loaded")),
            },
            "risk_policy": RISK,
            "app_registry": {
                "schema": APP_REGISTRY.get("schema") if isinstance(APP_REGISTRY, dict) else "",
                "app_count": len(APP_REGISTRY.get("apps", [])) if isinstance(APP_REGISTRY, dict) else 0,
            },
            "app_action_count": len(APP_ACTIONS),
            "base_action_count": len(ACTIONS) - len(APP_ACTIONS),
            "detail_hint": "Use system.action_schema for one action, system.app_registry for app details, or include_actions=true only when the full registry is explicitly required.",
        }

    def _action_life_activity_query(
        self,
        op_id: str,
        target: Optional[str],
        args: Dict[str, Any],
    ) -> Dict[str, Any]:
        del op_id, target
        provider = _LIFE_ACTIVITY_QUERY_PROVIDER
        if not callable(provider):
            raise OmniBodyError("authoritative Life activity query provider is unavailable")
        request = {
            "relative_day": str(args.get("relative_day") or "").strip(),
            "date": str(args.get("date") or "").strip(),
            "status": str(args.get("status") or "").strip(),
            "limit": args.get("limit", 30),
        }
        result = provider(request)
        if not isinstance(result, dict):
            raise OmniBodyError("Life activity query provider returned a non-object")
        if result.get("ok") is not True:
            raise OmniBodyError(str(result.get("error") or "Life activity query failed"))
        return result

    def _action_life_body_state_query(
        self,
        op_id: str,
        target: Optional[str],
        args: Dict[str, Any],
    ) -> Dict[str, Any]:
        del op_id, target
        provider = _BODY_STATE_QUERY_PROVIDER
        if not callable(provider):
            raise OmniBodyError("authoritative body state query provider is unavailable")
        request = {
            "sections": args.get("sections"),
            "recent_limit": args.get("recent_limit", 12),
        }
        result = provider(request)
        if not isinstance(result, dict):
            raise OmniBodyError("body state query provider returned a non-object")
        if result.get("ok") is not True:
            raise OmniBodyError(str(result.get("error") or "body state query failed"))
        return result

    def _action_system_app_registry(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        category = str(args.get("category") or "").strip()
        implemented_only = bool(args.get("implemented_only", False))
        apps = list(APP_REGISTRY.get("apps", [])) if isinstance(APP_REGISTRY, dict) else []
        if category:
            apps = [a for a in apps if a.get("category") == category]
        if implemented_only:
            apps = [a for a in apps if int(a.get("implemented_actions") or 0) > 0]
        return {
            "schema": APP_REGISTRY.get("schema", "tiangong.v3.omni_body.app_registry.v1") if isinstance(APP_REGISTRY, dict) else "tiangong.v3.omni_body.app_registry.v1",
            "total_apps": len(apps),
            "apps": apps,
        }

    def _action_system_action_schema(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        name = str(args.get("action") or target or "").strip()
        if not name:
            raise OmniBodyError("system.action_schema requires target or args.action")
        meta = ACTIONS.get(name)
        if meta is None:
            raise OmniBodyError(f"Unknown action: {name}")
        capability = self.capability_manifest.capabilities.get(name)
        return {"action": name, "schema": meta, "argument_contract": schema_for_action(name), "implemented": bool(meta.get("implemented")), "executable": bool(capability.executable) if capability else False, "capability": capability.to_dict() if capability else {}, "alias_to": meta.get("alias_to"), "adapter": meta.get("adapter")}

    def _action_system_health(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        deps = {}
        for mod in ["docx", "pptx", "openpyxl", "PIL", "pypdf", "reportlab"]:
            try:
                __import__(mod)
                deps[mod] = True
            except Exception:
                deps[mod] = False
        healthy = bool(self.capability_integrity.get("healthy")) and bool(deps.get("pptx"))
        return {
            "healthy": healthy,
            "workspace": str(self.workspace),
            "audit_dir": str(self.audit_dir),
            "backup_dir": str(self.backup_dir),
            "trash_dir": str(self.trash_dir),
            "config": asdict(self.config),
            "dependencies": deps,
            "ffmpeg": self.ffmpeg,
            "ffprobe": self.ffprobe,
            "app_registry_loaded": bool(APP_REGISTRY.get("apps")) if isinstance(APP_REGISTRY, dict) else False,
            "app_count": len(APP_REGISTRY.get("apps", [])) if isinstance(APP_REGISTRY, dict) else 0,
            "app_action_count": len(APP_ACTIONS),
            "required_capability_integrity": self.capability_integrity,
        }

    # ---------- learning ----------

    def _action_learning_ingest(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        expected_token = str(os.environ.get("TIANGONG_LEARNING_INGEST_TOKEN") or "").strip()
        provided_token = str(args.get("host_verified_intent_token") or "").strip()
        host_verified = bool(self.config.host_verified_learning_intent)
        legacy_token_verified = bool(expected_token and provided_token == expected_token)
        if not host_verified and not legacy_token_verified:
            return {
                "success": False,
                "ok": False,
                "op_id": op_id,
                "status": "blocked",
                "error": "host_verified_learning_intent_required",
                "message": "learning.ingest requires a host-verified intent token; use the server learning cards/from-request API when handling raw user dialogue.",
                "allowed_effect": "none",
                "evidence": {},
            }
        user_text = str(
            args.get("user_text")
            or args.get("userText")
            or args.get("instruction")
            or args.get("xiaoxi")
            or ""
        )
        material_text = args.get("material_text")
        if material_text is None:
            material_text = args.get("content")
        material_path = args.get("material_path")
        if material_path is None:
            material_path = args.get("path") or target
        provider = _LEARNING_INGEST_PROVIDER
        if callable(provider):
            result = provider({
                "user_text": user_text,
                "material_text": str(material_text) if material_text is not None else None,
                "material_path": str(material_path) if material_path else None,
                "source": str(args.get("source") or "omni_body.learning.ingest"),
                "desired_scope": str(args.get("desired_scope") or args.get("scope") or "skill"),
                "allow_network": bool(args.get("allow_network") is True),
                "actor": str(args.get("actor") or "model_tool"),
            })
            if not isinstance(result, dict):
                raise OmniBodyError("authoritative learning provider returned an invalid response")
            learning = result.get("learning") if isinstance(result.get("learning"), dict) else {}
            card_id = str(
                learning.get("learning_id")
                or result.get("card_id")
                or result.get("learning_id")
                or ""
            )
            status = str(learning.get("status") or result.get("status") or "")
            success = bool(result.get("ok")) and status == "awaiting_user" and bool(card_id)
            compact_result = {
                "ok": bool(result.get("ok")),
                "duplicate": bool(result.get("duplicate")),
                "card_id": card_id,
                "learning": {
                    "learning_id": card_id,
                    "status": status,
                    "target": str(learning.get("target") or ""),
                    "title": str(learning.get("title") or ""),
                    "risk_level": str(learning.get("risk_level") or ""),
                    "registered": bool(learning.get("registered")),
                    "requires_confirmation": bool(learning.get("requires_confirmation")),
                    "draft_sha256": str(learning.get("draft_sha256") or ""),
                },
            }
            return {
                "success": success,
                "ok": success,
                "op_id": op_id,
                "status": status,
                "card_id": card_id,
                "allowed_effect": "create_pending_learning_card_only",
                "forbidden_effects": ["compile_skill", "activate_skill", "register_tool", "release_tool"],
                # Do not return the full activity scope, materialization event,
                # and built preview to the chat model. They can exceed 200 KB
                # and bury the only fields needed for the next step.
                "result": compact_result,
                "evidence": {
                    "card_id": card_id,
                    "status": status,
                    "authority": "life_kernel",
                    "registered": bool(learning.get("registered")),
                    "exists": bool(card_id),
                },
                "message": (
                    "pending learning card created"
                    if success
                    else str(result.get("error") or result.get("reason_code") or status or "learning card creation failed")
                ),
            }
        Engine, default_root, import_error = _load_learning_runtime()
        if Engine is None:
            raise OmniBodyError(f"learning runtime unavailable: {import_error}")
        root = default_root
        engine = Engine(root=root)
        result = engine.create_learning_card_from_request(
            user_text=user_text,
            material_text=str(material_text) if material_text is not None else None,
            material_path=str(material_path) if material_path else None,
            source=str(args.get("source") or "omni_body.learning.ingest"),
            desired_scope=str(args.get("desired_scope") or args.get("scope") or "skill"),
            allow_network=bool(args.get("allow_network") is True),
            actor=str(args.get("actor") or "model_tool"),
        )
        success = bool(result.get("ok"))
        return {
            "success": success,
            "ok": success,
            "op_id": op_id,
            "status": result.get("status"),
            "card_id": result.get("card_id"),
            "allowed_effect": "create_pending_learning_card_only",
            "forbidden_effects": ["compile_skill", "activate_skill", "register_tool", "release_tool"],
            "result": result,
            "evidence": {
                "card_id": result.get("card_id"),
                "status": result.get("status"),
                "learning_root": str(root),
                "path": str(root / "cards" / f"{result.get('card_id')}.json") if result.get("card_id") else "",
                "exists": bool(result.get("card_id")),
            },
            "message": result.get("error") or result.get("status") or "learning card created",
        }

    # ---------- files ----------

    def _action_file_list(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        root = self._resolve(target, must_exist=True)
        pattern = args.get("pattern", "*")
        recursive = bool(args.get("recursive", False))
        include_hidden = bool(args.get("include_hidden", False))
        max_results = int(args.get("max_results", 500))
        iterator = root.rglob(pattern) if recursive else root.glob(pattern)
        entries = []
        for p in iterator:
            if not include_hidden and any(part.startswith(".") for part in p.relative_to(root).parts):
                continue
            entries.append({
                "name": p.name,
                "path": str(p),
                "rel_path": self._rel(p),
                "type": "dir" if p.is_dir() else "file",
                "size_bytes": p.stat().st_size if p.is_file() else None,
                "modified": p.stat().st_mtime,
            })
            if len(entries) >= max_results:
                break
        return {"root": str(root), "count": len(entries), "entries": entries}

    @staticmethod
    def _canonical_text_encoding(args: Dict[str, Any]) -> str:
        raw = str(args.get("encoding") or "utf-8").strip().lower().replace("_", "-")
        if raw not in {"utf-8", "utf8"}:
            raise OmniBodyError("text operations require canonical UTF-8 encoding")
        return "utf-8"

    @staticmethod
    def _read_canonical_utf8(path: Path) -> str:
        raw = path.read_bytes()
        if raw.startswith((b"\xef\xbb\xbf", b"\xff\xfe", b"\xfe\xff", b"\x00\x00\xfe\xff", b"\xff\xfe\x00\x00")):
            raise OmniBodyError("text files must use UTF-8 without BOM")
        try:
            return raw.decode("utf-8", errors="strict")
        except UnicodeDecodeError as exc:
            raise OmniBodyError("text file is not valid UTF-8") from exc

    @staticmethod
    def _write_canonical_utf8(path: Path, content: Any) -> None:
        try:
            encoded = str(content).encode("utf-8", errors="strict")
        except UnicodeEncodeError as exc:
            raise OmniBodyError("text content contains an invalid Unicode scalar") from exc
        # Binary write deliberately bypasses platform newline translation so
        # Linux and Windows produce identical bytes and content hashes.
        path.write_bytes(encoded)

    def _action_file_read(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target, must_exist=True)
        if p.is_dir():
            raise OmniBodyError("file.read target is a directory; use file.list")
        binary = bool(args.get("binary", False))
        max_chars = int(args.get("max_chars", self.config.max_text_read_chars))
        if binary:
            raw = p.read_bytes()
            preview = base64.b64encode(raw[:max_chars]).decode("ascii")
            return {"path": str(p), "mime": mimetypes.guess_type(str(p))[0], "size_bytes": len(raw), "base64_preview": preview, "truncated": len(raw) > max_chars}
        self._canonical_text_encoding(args)
        data = self._read_canonical_utf8(p)
        return {"path": str(p), "size_chars": len(data), "content": data[:max_chars], "truncated": len(data) > max_chars, "evidence": self._file_evidence(p)}

    def _action_file_write(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target)
        snapshots = self._snapshot(op_id, [p])
        p.parent.mkdir(parents=True, exist_ok=True)
        binary = bool(args.get("binary", False))
        if binary:
            content_b64 = args.get("base64") or args.get("content")
            if content_b64 is None:
                raise OmniBodyError("binary file.write requires args.base64 or args.content")
            p.write_bytes(base64.b64decode(content_b64))
        else:
            content = args.get("content", "")
            self._canonical_text_encoding(args)
            self._write_canonical_utf8(p, content)
        return {"snapshots": snapshots, "evidence": self._file_evidence(p)}

    def _action_file_append(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target)
        snapshots = self._snapshot(op_id, [p])
        p.parent.mkdir(parents=True, exist_ok=True)
        self._canonical_text_encoding(args)
        try:
            encoded = str(args.get("content", "")).encode("utf-8", errors="strict")
        except UnicodeEncodeError as exc:
            raise OmniBodyError("text content contains an invalid Unicode scalar") from exc
        with p.open("ab") as f:
            f.write(encoded)
        return {"snapshots": snapshots, "evidence": self._file_evidence(p)}

    def _copy_path(self, src: Path, dst: Path, overwrite: bool = False) -> None:
        if dst.exists():
            if not overwrite:
                raise FileExistsError(str(dst))
            self._remove_path(dst)
        dst.parent.mkdir(parents=True, exist_ok=True)
        if src.is_dir():
            shutil.copytree(src, dst)
        else:
            shutil.copy2(src, dst)

    def _action_file_copy(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        src = self._resolve(target, must_exist=True)
        dst = self._resolve(args.get("destination") or args.get("target"))
        snapshots = self._snapshot(op_id, [dst])
        self._copy_path(src, dst, overwrite=bool(args.get("overwrite", False)))
        return {"snapshots": snapshots, "source": self._file_evidence(src), "destination": self._file_evidence(dst)}

    def _action_file_move(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        src = self._resolve(target, must_exist=True)
        dst = self._resolve(args.get("destination") or args.get("target"))
        snapshots = self._snapshot(op_id, [src, dst])
        if dst.exists():
            if not args.get("overwrite", False):
                raise FileExistsError(str(dst))
            self._remove_path(dst)
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src), str(dst))
        return {"snapshots": snapshots, "source_exists_after": src.exists(), "destination": self._file_evidence(dst, with_hash=False)}

    def _action_file_rename(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        src = self._resolve(target, must_exist=True)
        if src in {self.workspace, self.audit_dir, self.backup_dir, self.trash_dir}:
            raise OmniBodyError("Refusing to rename a protected runtime root")
        new_name = _validate_cross_platform_filename(args.get("new_name"))
        # Keep the invocation inside the signed relative workspace contract.
        # Passing the already-resolved absolute path back through file.move
        # incorrectly asks the A4 gate to authorize a new absolute-path call.
        src_relative = self._rel(src)
        dst_relative = (Path(src_relative).parent / new_name).as_posix()
        return self._action_file_move(
            op_id,
            src_relative,
            {"destination": dst_relative, "overwrite": bool(args.get("overwrite", False))},
        )

    def _action_file_mkdir(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target)
        snapshots = self._snapshot(op_id, [p])
        p.mkdir(parents=True, exist_ok=bool(args.get("exist_ok", True)))
        return {"snapshots": snapshots, "evidence": self._file_evidence(p)}

    def _action_file_delete_to_trash(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target, must_exist=True)
        if p in [self.workspace, self.audit_dir, self.backup_dir, self.trash_dir]:
            raise OmniBodyError("Refusing to trash workspace/audit/backup/trash root")
        snapshots = self._snapshot(op_id, [p])
        trash_target = self.trash_dir / op_id / self._rel(p)
        trash_target.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(p), str(trash_target))
        return {"snapshots": snapshots, "trash_path": str(trash_target), "source_exists_after": p.exists()}

    def _action_file_search(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        root = self._resolve(target or ".", must_exist=True)
        pattern = args.get("pattern", "*")
        text = args.get("text")
        use_regex = bool(args.get("regex", False))
        recursive = bool(args.get("recursive", True))
        max_results = int(args.get("max_results", 100))
        max_size = self.config.max_search_file_size_mb * 1024 * 1024
        matches = []
        iterator = root.rglob(pattern) if recursive else root.glob(pattern)
        regex = re.compile(text, re.I) if text and use_regex else None
        for p in iterator:
            if not p.is_file():
                continue
            hit = False
            snippet = None
            if text is None:
                hit = True
            elif p.stat().st_size <= max_size:
                data = p.read_text(encoding=args.get("encoding", "utf-8"), errors="ignore")
                if regex:
                    m = regex.search(data)
                    if m:
                        hit = True
                        snippet = data[max(0, m.start()-80):m.end()+80]
                else:
                    idx = data.lower().find(str(text).lower())
                    if idx >= 0:
                        hit = True
                        snippet = data[max(0, idx-80):idx+len(str(text))+80]
            if hit:
                matches.append({"path": str(p), "rel_path": self._rel(p), "size_bytes": p.stat().st_size, "snippet": snippet})
                if len(matches) >= max_results:
                    break
        return {"root": str(root), "count": len(matches), "matches": matches}

    def _action_file_hash(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target, must_exist=True)
        if not p.is_file():
            raise OmniBodyError("file.hash requires a file")
        return {"path": str(p), "sha256": self._sha256(p), "size_bytes": p.stat().st_size}

    # ---------- zip ----------

    def _action_zip_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        output = self._resolve(target or args.get("output") or "archive.zip")
        sources = args.get("sources") or []
        if isinstance(sources, str):
            sources = [sources]
        if not sources:
            raise OmniBodyError("zip.create requires args.sources")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for src_str in sources:
                src = self._resolve(src_str, must_exist=True)
                if src.is_dir():
                    for p in src.rglob("*"):
                        if p.is_file():
                            zf.write(p, arcname=p.relative_to(src.parent))
                else:
                    zf.write(src, arcname=src.name)
        return {"snapshots": snapshots, "output": self._file_evidence(output)}

    def _action_zip_extract(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        zpath = self._resolve(target, must_exist=True)
        dest = self._resolve(args.get("destination") or zpath.with_suffix("" ).name)
        snapshots = self._snapshot(op_id, [dest])
        dest.mkdir(parents=True, exist_ok=True)
        extracted = []
        with zipfile.ZipFile(zpath, "r") as zf:
            for member in zf.infolist():
                out_path = (dest / member.filename).resolve()
                if not self._is_inside(out_path, dest):
                    raise OmniBodyError(f"Unsafe zip member path: {member.filename}")
            zf.extractall(dest)
            extracted = [str((dest / m.filename).resolve()) for m in zf.infolist() if not m.is_dir()]
        return {"snapshots": snapshots, "destination": str(dest), "extracted_count": len(extracted), "extracted_preview": extracted[:100]}

    # ---------- code / quality / execution ----------

    def _action_code_read(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        res = self._action_file_read(op_id, target, args)
        p = self._resolve(target, must_exist=True)
        res["language"] = args.get("language") or self._language_from_suffix(p.suffix)
        return res

    def _language_from_suffix(self, suffix: str) -> str:
        return {
            ".py": "python", ".js": "javascript", ".ts": "typescript", ".tsx": "typescript-react", ".jsx": "javascript-react",
            ".java": "java", ".go": "go", ".rs": "rust", ".cpp": "cpp", ".c": "c", ".cs": "csharp", ".php": "php",
            ".md": "markdown", ".json": "json", ".yaml": "yaml", ".yml": "yaml", ".html": "html", ".css": "css",
        }.get(suffix.lower(), "text")

    def _action_code_write(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        res = self._action_file_write(op_id, target, args)
        p = self._resolve(target, must_exist=True)
        language = args.get("language") or self._language_from_suffix(p.suffix)
        checks = []
        if args.get("syntax_check", True) and language == "python":
            checks.append(self._compile_python_file(p))
        if args.get("syntax_check", True) and language in {"javascript", "javascript-react"} and p.suffix.lower() in {".js", ".mjs", ".cjs"}:
            checks.append(self._compile_javascript_file(p))
        res["language"] = language
        res["quality_checks"] = checks
        return res

    def _compile_python_file(self, p: Path) -> Dict[str, Any]:
        try:
            compile(p.read_text(encoding="utf-8"), str(p), "exec")
            return {"path": str(p), "ok": True, "type": "python_syntax"}
        except SyntaxError as e:
            return {"path": str(p), "ok": False, "type": "python_syntax", "line": e.lineno, "offset": e.offset, "message": e.msg}

    def _action_code_patch_replace(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target, must_exist=True)
        if not p.is_file():
            raise OmniBodyError("code.patch_replace target must be an existing file")
        find = args.get("find")
        replace = args.get("replace", "")
        if not isinstance(find, str) or not find:
            raise OmniBodyError("code.patch_replace requires non-empty args.find")
        if not isinstance(replace, str):
            raise OmniBodyError("code.patch_replace args.replace must be a string")
        count = int(args.get("count", 0))
        if count < 0:
            raise OmniBodyError("code.patch_replace args.count must be non-negative")
        self._canonical_text_encoding(args)
        data = self._read_canonical_utf8(p)
        if args.get("regex", False):
            new_data, n = re.subn(find, replace, data, count=count if count > 0 else 0)
        else:
            n = data.count(find) if count <= 0 else min(count, data.count(find))
            new_data = data.replace(find, replace, count if count > 0 else -1)
        if n == 0 and not args.get("allow_noop", False):
            raise OmniBodyError("No replacement occurred")
        # Snapshot only the validated target file, after every parameter and
        # no-op check has passed but before the write begins.
        snapshots = self._snapshot(op_id, [p])
        self._write_canonical_utf8(p, new_data)
        return {"snapshots": snapshots, "replacements": n, "evidence": self._file_evidence(p)}

    def _compile_javascript_file(self, p: Path) -> Dict[str, Any]:
        node = shutil.which("node") or shutil.which("node.exe")
        if not node:
            return {"path": str(p), "ok": False, "type": "javascript_syntax", "message": "node executable is unavailable"}
        result = self._run_subprocess([node, "--check", str(p)], timeout=30, cwd=self.workspace)
        return {
            "path": str(p),
            "ok": bool(result.get("ok")),
            "type": "javascript_syntax",
            "returncode": result.get("returncode"),
            "stdout": result.get("stdout", ""),
            "stderr": result.get("stderr", ""),
        }

    def _action_quality_python_syntax(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        root = self._resolve(target or ".", must_exist=True)
        files = [root] if root.is_file() else list(root.rglob("*.py"))
        checks = [self._compile_python_file(p) for p in files]
        ok = all(c.get("ok") for c in checks)
        return {"ok": ok, "checked_count": len(checks), "checks": checks, "success": ok}

    def _action_quality_javascript_syntax(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        root = self._resolve(target, must_exist=True)
        suffixes = {".js", ".mjs", ".cjs"}
        if root.is_file():
            files = [root] if root.suffix.lower() in suffixes else []
        else:
            iterator = root.rglob("*") if args.get("recursive", True) else root.glob("*")
            files = [path for path in iterator if path.is_file() and path.suffix.lower() in suffixes]
        checks = [self._compile_javascript_file(path) for path in sorted(files)]
        ok = bool(checks) and all(check.get("ok") for check in checks)
        return {"ok": ok, "checked_count": len(checks), "checks": checks, "success": ok}

    def _action_quality_run_tests(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        if not self.config.allow_shell:
            raise OmniBodyError("quality.run_tests requires BodyRuntimeConfig.allow_shell=True")
        cmd = args.get("command") or [_resolve_python_interpreter(), "-m", "pytest", "-q"]
        cmd_list, _executable = self._normalize_shell_command(cmd, action="quality.run_tests")
        run_cwd = self._resolve(target or ".", must_exist=True)
        if run_cwd.is_file():
            run_cwd = run_cwd.parent
        res = self._run_subprocess(
            cmd_list,
            timeout=int(args.get("timeout", self.config.default_timeout_seconds)),
            shell=False,
            cwd=run_cwd,
            op_id=op_id,
        )
        return {"command": cmd_list, "cwd": str(run_cwd), "execution": res, "success": res["ok"]}

    def _action_python_run(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        if not self.config.allow_python:
            raise OmniBodyError("python.run requires BodyRuntimeConfig.allow_python=True")
        timeout = int(args.get("timeout", self.config.default_timeout_seconds))
        python_executable = _resolve_python_interpreter()
        if target:
            script = self._resolve(target, must_exist=True)
            cmd = [python_executable, str(script)] + list(args.get("argv", []))
            res = self._run_subprocess(cmd, timeout=timeout, op_id=op_id)
            return {"command": cmd, "execution": res, "success": res["ok"]}
        code = args.get("code")
        if code is None:
            raise OmniBodyError("python.run requires target script or args.code")
        with tempfile.NamedTemporaryFile("w", suffix=".py", dir=self.workspace, delete=False, encoding="utf-8") as f:
            f.write(code)
            temp_name = f.name
        try:
            res = self._run_subprocess([python_executable, temp_name], timeout=timeout, op_id=op_id)
            return {"temp_script": temp_name, "execution": res, "success": res["ok"]}
        finally:
            try:
                Path(temp_name).unlink()
            except FileNotFoundError:
                pass

    def _action_git_clone(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        """Governed GitHub network-read clone; never delegates to shell.run."""
        repo_url = str(target or args.get("url") or "").strip()
        destination = str(args.get("destination") or "").strip()
        if not repo_url:
            raise OmniBodyError("git.clone requires a GitHub HTTPS repository URL as target")
        if not destination:
            raise OmniBodyError("git.clone requires args.destination")
        output = self._resolve(destination, must_exist=False)
        if output.exists():
            raise OmniBodyError("git.clone destination must be a new directory")
        try:
            from .network_capabilities import NetworkCapabilityError, clone_public_github_repo
            result = clone_public_github_repo(
                repo_url,
                output,
                timeout_seconds=int(args.get("timeout", 300)),
            )
        except NetworkCapabilityError as exc:
            raise OmniBodyError(str(exc)) from exc
        return {
            **result,
            "operation_id": op_id,
            "network_authority": "typed:git.clone:github.com:https:read-only",
            "generic_shell_network": "denied",
        }

    def _action_shell_run(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        if not self.config.allow_shell:
            raise OmniBodyError("shell.run requires BodyRuntimeConfig.allow_shell=True")
        cmd = args.get("command")
        if not cmd:
            raise OmniBodyError("shell.run requires args.command")
        cmd_list, executable = self._normalize_shell_command(cmd, action="shell.run")
        if self.config.allowed_shell_commands:
            if executable not in self.config.allowed_shell_commands:
                raise OmniBodyError(f"Shell command not in allowlist: {executable}")
        res = self._run_subprocess(cmd_list, timeout=int(args.get("timeout", self.config.default_timeout_seconds)), shell=False, op_id=op_id)
        return {"command": cmd_list, "execution": res, "success": res["ok"]}

    def _action_word_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        """word.create 别名 → docx.create"""
        return self._action_docx_create(op_id, target, args)

    def _action_word_read(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        """Read semantic text and table content from a .docx file."""
        try:
            from docx import Document
        except Exception as exc:
            raise OmniBodyError("python-docx is required for word.read") from exc
        src = self._resolve(target or args.get("path"), must_exist=True)
        if src.suffix.lower() != ".docx":
            raise OmniBodyError("word.read supports .docx files")
        doc = Document(src)
        max_chars = max(1, int(args.get("max_chars", self.config.max_text_read_chars)))
        paragraphs = [p.text for p in doc.paragraphs if str(p.text or "").strip()]
        tables: List[List[List[str]]] = []
        table_lines: List[str] = []
        for table in doc.tables:
            rows: List[List[str]] = []
            for row in table.rows:
                values = [str(cell.text or "") for cell in row.cells]
                rows.append(values)
                table_lines.append(" | ".join(values))
            tables.append(rows)
        text = "\n".join([*paragraphs, *table_lines])
        truncated = len(text) > max_chars
        return {
            "path": str(src),
            "text": text[:max_chars],
            "paragraphs": paragraphs,
            "tables": tables,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "truncated": truncated,
            "success": True,
        }

    def _action_docx_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        try:
            from docx import Document
            from docx.shared import Cm, Pt
            from docx.oxml.ns import qn
        except Exception as e:
            raise OmniBodyError("python-docx is required for docx.create") from e
        output = self._resolve(target or args.get("output") or "document.docx")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        # MM-P2-GROUP-1: explicit Chinese document defaults — visible fonts,
        # standard margins and heading east-asian fonts (no bare Document()).
        try:
            normal = doc.styles["Normal"]
            normal.font.name = "Times New Roman"
            normal.font.size = Pt(12)
            normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
            for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
                try:
                    style = doc.styles[style_name]
                    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
                except Exception:
                    pass
            for section in doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.17)
                section.right_margin = Cm(3.17)
        except Exception:
            pass

        # —— 纯文本 content 直出模式（模型最常用的传参方式）——
        content = args.get("content") or args.get("text") or ""
        source = args.get("source") or args.get("source_path") or args.get("markdown_path") or args.get("markdown_file")
        source_path = None
        if not (isinstance(content, str) and content.strip()) and isinstance(source, str) and source.strip():
            source_path = self._resolve(source, must_exist=True)
            if source_path.suffix.lower() not in {".md", ".txt"} or not source_path.is_file():
                raise OmniBodyError("docx.create source must be an existing workspace .md or .txt file")
            if source_path.stat().st_size > 8 * 1024 * 1024:
                raise OmniBodyError("docx.create source exceeds 8 MiB")
            content = source_path.read_text(encoding="utf-8")
        if isinstance(content, str) and content.strip():
            # 按 markdown 风格拆 # / ## / 空行 / |表格| / -列表
            for block in content.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                if block.startswith("# ") or block.startswith("## ") or block.startswith("### "):
                    level = block.count("#", 0, 4)
                    doc.add_heading(block.lstrip("# ").strip(), level=min(level, 3))
                elif block.startswith("|") and block.endswith("|"):
                    lines = [ln for ln in block.split("\n") if ln.strip().startswith("|")]
                    if len(lines) >= 2:
                        headers = [c.strip() for c in lines[0].strip("|").split("|")]
                        rows = [[c.strip() for c in ln.strip("|").split("|")] for ln in lines[2:]]
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = "Table Grid"
                        for i, h in enumerate(headers):
                            table.rows[0].cells[i].text = h
                        for row in rows:
                            cells = table.add_row().cells
                            for i, val in enumerate(row[:len(cells)]):
                                cells[i].text = val
                elif block.startswith("- ") or block.startswith("* "):
                    for line in block.split("\n"):
                        doc.add_paragraph(line.lstrip("- *").strip(), style="List Bullet")
                else:
                    doc.add_paragraph(block)
            doc.save(output)
            return {"snapshots": snapshots, "output": self._file_evidence(output), "content_mode": "source_text" if source_path else "flat_text", "source": str(source_path) if source_path else "", "chars": len(content)}

        # —— 结构化 sections 模式 ——
        has_content = False
        if args.get("title"):
            doc.add_heading(str(args["title"]), 0)
            has_content = True
        if args.get("subtitle"):
            p = doc.add_paragraph(str(args["subtitle"]))
            p.style = doc.styles["Subtitle"] if "Subtitle" in [s.name for s in doc.styles] else p.style
            has_content = True
        for section in args.get("sections", []):
            if section.get("heading"):
                doc.add_heading(str(section["heading"]), int(section.get("level", 1)))
                has_content = True
            for para in section.get("paragraphs", []):
                doc.add_paragraph(str(para))
                has_content = True
            for bullet in section.get("bullets", []):
                doc.add_paragraph(str(bullet), style="List Bullet")
                has_content = True
            table_spec = section.get("table")
            if table_spec:
                headers = table_spec.get("headers", [])
                rows = table_spec.get("rows", [])
                table = doc.add_table(rows=1 if headers else 0, cols=max(1, len(headers) or (len(rows[0]) if rows else 1)))
                table.style = table_spec.get("style", "Table Grid")
                if headers:
                    for i, h in enumerate(headers):
                        table.rows[0].cells[i].text = str(h)
                for row in rows:
                    cells = table.add_row().cells
                    for i, val in enumerate(row[:len(cells)]):
                        cells[i].text = str(val)
                has_content = True

        if not has_content:
            return {"ok": False, "error": "docx.create 需要 content 或 title/sections 参数，不能生成空文档", "args_received": {k: str(v)[:80] for k, v in args.items()}}

        doc.save(output)
        return {"snapshots": snapshots, "output": self._file_evidence(output), "sections_count": len(args.get("sections", []))}

    def _action_pptx_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        try:
            from .ppt_design import build_presentation
        except Exception as e:
            raise OmniBodyError(f"PPT design engine is unavailable: {type(e).__name__}: {e}") from e
        output = self._resolve(target or args.get("output") or "deck.pptx")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        content = args.get("content") or args.get("text") or ""
        source = args.get("source") or args.get("source_path") or args.get("markdown_path") or args.get("markdown_file")
        source_path = None
        if not (isinstance(content, str) and content.strip()) and isinstance(source, str) and source.strip():
            source_path = self._resolve(source, must_exist=True)
            if source_path.suffix.lower() not in {".md", ".txt"} or not source_path.is_file():
                raise OmniBodyError("pptx.create source must be an existing workspace .md or .txt file")
            if source_path.stat().st_size > 4 * 1024 * 1024:
                raise OmniBodyError("pptx.create source exceeds 4 MiB")
            content = source_path.read_text(encoding="utf-8")
        try:
            prs, design_evidence = build_presentation(
                self._resolve,
                args,
                content=str(content or ""),
                source_path=source_path,
            )
        except (OSError, UnicodeError, ValueError, json.JSONDecodeError) as exc:
            raise OmniBodyError(f"pptx.create design/content validation failed: {exc}") from exc
        prs.save(output)
        return {
            "snapshots": snapshots,
            "output": self._file_evidence(output),
            "content_mode": "source_text" if source_path else ("flat_text" if str(content or "").strip() else "structured"),
            "source": str(source_path) if source_path else "",
            "slides_count": len(prs.slides),
            "design": design_evidence,
        }

    def _action_pptx_read(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        try:
            from .ppt_design import inspect_presentation
        except Exception as exc:
            raise OmniBodyError(f"PPT inspection engine is unavailable: {type(exc).__name__}: {exc}") from exc
        source = self._resolve(target or args.get("path"), must_exist=True)
        if not source.is_file() or source.suffix.lower() != ".pptx":
            raise OmniBodyError("pptx.read target must be an existing .pptx file")
        inspection = inspect_presentation(source)
        max_chars = max(200, min(int(args.get("max_chars_per_slide", 2400)), 20_000))
        for slide in inspection.get("slides", []):
            text = str(slide.get("text") or "")
            slide["text"] = text[:max_chars]
            slide["text_truncated"] = len(text) > max_chars
        return {
            "success": True,
            "path": str(source),
            "inspection": inspection,
            "evidence": self._file_evidence(source),
        }

    def _action_sheet_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        output = self._resolve(target or args.get("output") or "workbook.xlsx")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)

        # —— 纯文本 content 直出模式 ——
        content = args.get("content") or args.get("text") or ""
        if isinstance(content, str) and content.strip():
            # 按 |表格| markdown 或 tab 分隔解析
            rows = []
            for line in content.split("\n"):
                line = line.strip()
                if not line or line.startswith("|-") or line.startswith("|--"):
                    continue
                if line.startswith("|") and line.endswith("|"):
                    rows.append([c.strip() for c in line.strip("|").split("|")])
                elif "\t" in line:
                    rows.append(line.split("\t"))
                elif "," in line and not line.startswith("#"):
                    rows.append([c.strip() for c in line.split(",")])
            if rows:
                headers = rows[0] if rows else []
                data = rows[1:] if len(rows) > 1 else []
                sheets = [{"name": "Sheet1", "headers": headers, "rows": data}]
            else:
                sheets = [{"name": "Sheet1", "headers": [], "rows": [[ln] for ln in content.split("\n") if ln.strip()]}]
        else:
            sheets = args.get("sheets") or [{"name": "Sheet1", "headers": args.get("headers", []), "rows": args.get("rows", [])}]

        if output.suffix.lower() == ".csv":
            first = sheets[0] if sheets else {"headers": [], "rows": []}
            rows = []
            if first.get("headers"):
                rows.append(first.get("headers"))
            rows.extend(first.get("rows", []))
            self._canonical_text_encoding(args)
            with output.open("w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                writer.writerows(rows)
            return {"snapshots": snapshots, "output": self._file_evidence(output), "sheets_count": 1, "writer": "csv"}

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
        except Exception:
            self._write_minimal_xlsx(output, sheets)
            return {"snapshots": snapshots, "output": self._file_evidence(output), "sheets_count": len(sheets), "writer": "stdlib_minimal_xlsx"}

        wb = Workbook()
        default = wb.active
        try:
            for idx, spec in enumerate(sheets):
                ws = default if idx == 0 else wb.create_sheet()
                ws.title = str(spec.get("name", f"Sheet{idx+1}"))[:31]
                rows = []
                if spec.get("headers"):
                    rows.append(spec.get("headers"))
                rows.extend(spec.get("rows", []))
                for row in rows:
                    ws.append(list(row))
                if spec.get("headers"):
                    for cell in ws[1]:
                        cell.font = Font(bold=True)
                        cell.fill = PatternFill("solid", fgColor="D9EAF7")
                        cell.alignment = Alignment(horizontal="center")
                for col in ws.columns:
                    max_len = max(len(str(c.value)) if c.value is not None else 0 for c in col)
                    ws.column_dimensions[col[0].column_letter].width = min(max(max_len + 2, 10), 40)
            wb.save(output)
        finally:
            try:
                wb.close()
            except Exception:
                pass
        return {"snapshots": snapshots, "output": self._file_evidence(output), "sheets_count": len(sheets), "writer": "openpyxl"}

    def _action_sheet_read(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target, must_exist=True)
        max_rows = int(args.get("max_rows", 20))
        if p.suffix.lower() == ".csv":
            self._canonical_text_encoding(args)
            with p.open("r", encoding="utf-8", errors="strict", newline="") as f:
                rows = []
                for i, row in enumerate(csv.reader(f)):
                    rows.append(row)
                    if i + 1 >= max_rows:
                        break
            return {"path": str(p), "type": "csv", "rows_preview": rows}
        try:
            from openpyxl import load_workbook
            wb = load_workbook(p, data_only=bool(args.get("data_only", False)), read_only=True)
            try:
                previews = {}
                for ws in wb.worksheets[: int(args.get("max_sheets", 5))]:
                    rows = []
                    for row in ws.iter_rows(max_row=max_rows, values_only=True):
                        rows.append(list(row))
                    previews[ws.title] = rows
                sheetnames = list(wb.sheetnames)
            finally:
                wb.close()
            return {"path": str(p), "type": "xlsx", "sheets": sheetnames, "previews": previews, "reader": "openpyxl"}
        except ModuleNotFoundError:
            return self._read_minimal_xlsx(p, max_rows=max_rows, max_sheets=int(args.get("max_sheets", 5)))

    def _xlsx_col(self, n: int) -> str:
        s = ""
        while n:
            n, r = divmod(n - 1, 26)
            s = chr(65 + r) + s
        return s or "A"

    def _safe_sheet_name(self, name: Any, index: int) -> str:
        raw = str(name or f"Sheet{index}")
        return re.sub(r"[\\/*?:\[\]]", "_", raw)[:31] or f"Sheet{index}"

    def _write_minimal_xlsx(self, output: Path, sheets: List[Dict[str, Any]]) -> None:
        import html
        from datetime import datetime, timezone

        def esc(value: Any) -> str:
            return html.escape("" if value is None else str(value), quote=False)

        sheets = sheets or [{"name": "Sheet1", "headers": [], "rows": []}]
        now = datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
        content_types = [
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">',
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
            '<Default Extension="xml" ContentType="application/xml"/>',
            '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>',
            '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>',
            '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>',
        ]
        for idx, _ in enumerate(sheets, 1):
            content_types.append(f'<Override PartName="/xl/worksheets/sheet{idx}.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>')
        content_types.append('</Types>')

        workbook_sheets = []
        workbook_rels = []
        sheet_xmls: Dict[str, str] = {}
        for idx, spec in enumerate(sheets, 1):
            name = self._safe_sheet_name(spec.get("name"), idx)
            workbook_sheets.append(f'<sheet name="{html.escape(name, quote=True)}" sheetId="{idx}" r:id="rId{idx}"/>')
            workbook_rels.append(f'<Relationship Id="rId{idx}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet{idx}.xml"/>')
            rows = []
            if spec.get("headers"):
                rows.append(spec.get("headers"))
            rows.extend(spec.get("rows", []))
            row_xml = []
            for r_idx, row in enumerate(rows, 1):
                cells = []
                for c_idx, value in enumerate(list(row), 1):
                    ref = f"{self._xlsx_col(c_idx)}{r_idx}"
                    cells.append(f'<c r="{ref}" t="inlineStr"><is><t>{esc(value)}</t></is></c>')
                row_xml.append(f'<row r="{r_idx}">{"".join(cells)}</row>')
            sheet_xmls[f"xl/worksheets/sheet{idx}.xml"] = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
                f'<sheetData>{"".join(row_xml)}</sheetData></worksheet>'
            )

        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", "".join(content_types))
            zf.writestr("_rels/.rels", '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/><Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/></Relationships>')
            zf.writestr("docProps/core.xml", f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:creator>tiangong omni_body</dc:creator><cp:lastModifiedBy>tiangong omni_body</cp:lastModifiedBy><dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created><dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified></cp:coreProperties>')
            zf.writestr("docProps/app.xml", '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"><Application>tiangong omni_body</Application></Properties>')
            zf.writestr("xl/workbook.xml", '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><sheets>' + "".join(workbook_sheets) + '</sheets></workbook>')
            zf.writestr("xl/_rels/workbook.xml.rels", '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">' + "".join(workbook_rels) + '</Relationships>')
            for path, xml in sheet_xmls.items():
                zf.writestr(path, xml)

    def _read_minimal_xlsx(self, p: Path, max_rows: int = 20, max_sheets: int = 5) -> Dict[str, Any]:
        import xml.etree.ElementTree as ET
        ns = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main", "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
        rel_ns = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
        previews: Dict[str, Any] = {}
        sheet_names: List[str] = []
        with zipfile.ZipFile(p, "r") as zf:
            workbook = ET.fromstring(zf.read("xl/workbook.xml"))
            rels = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
            rid_to_target = {rel.attrib["Id"]: rel.attrib["Target"] for rel in rels.findall("rel:Relationship", rel_ns)}
            shared_strings: List[str] = []
            if "xl/sharedStrings.xml" in zf.namelist():
                shared_root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
                for si in shared_root.findall("m:si", ns):
                    shared_strings.append("".join(t.text or "" for t in si.findall(".//m:t", ns)))
            for sheet in workbook.findall("m:sheets/m:sheet", ns)[:max_sheets]:
                name = sheet.attrib.get("name", "Sheet")
                rid = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                target = rid_to_target.get(rid or "")
                if not target:
                    continue
                sheet_names.append(name)
                sheet_path = "xl/" + target if not target.startswith("/") and not target.startswith("xl/") else target.lstrip("/")
                root = ET.fromstring(zf.read(sheet_path))
                rows = []
                for row in root.findall("m:sheetData/m:row", ns)[:max_rows]:
                    vals = []
                    for c in row.findall("m:c", ns):
                        cell_type = c.attrib.get("t")
                        if cell_type == "inlineStr":
                            vals.append("".join(t.text or "" for t in c.findall(".//m:t", ns)))
                        else:
                            v = c.find("m:v", ns)
                            raw = v.text if v is not None else ""
                            if cell_type == "s" and raw.isdigit() and int(raw) < len(shared_strings):
                                vals.append(shared_strings[int(raw)])
                            else:
                                vals.append(raw)
                    rows.append(vals)
                previews[name] = rows
        return {"path": str(p), "type": "xlsx", "sheets": sheet_names, "previews": previews, "reader": "stdlib_minimal_xlsx"}

    def _action_mindmap_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        requested_output = self._resolve(target or args.get("output") or "mindmap.md")
        requested_suffix = requested_output.suffix.lower()
        if requested_suffix == ".html":
            output = requested_output.with_suffix(".md")
            html_path = requested_output
        elif requested_suffix == ".opml":
            output = requested_output.with_suffix(".md")
            html_path = requested_output.with_suffix(".html")
        else:
            output = requested_output
            html_path = requested_output.with_suffix(".html")
        opml_path = (
            requested_output
            if requested_suffix == ".opml"
            else requested_output.with_suffix(".opml")
            if args.get("opml", False)
            else None
        )
        artifact_paths = [path for path in (output, html_path, opml_path) if path is not None]
        if len({os.path.normcase(str(path)) for path in artifact_paths}) != len(artifact_paths):
            raise OmniBodyError("mindmap output paths must be distinct")
        snapshots = self._snapshot(op_id, artifact_paths)
        title = args.get("title", "Mindmap")
        tree = args.get("tree") or args.get("outline") or []

        # —— 纯文本 content 直出模式 ——
        content = args.get("content") or args.get("text") or ""
        source = args.get("source") or args.get("source_path") or args.get("markdown_path") or args.get("markdown_file")
        source_path = None
        if not (isinstance(content, str) and content.strip()) and not tree and isinstance(source, str) and source.strip():
            source_path = self._resolve(source, must_exist=True)
            if source_path.suffix.lower() not in {".md", ".txt"} or not source_path.is_file():
                raise OmniBodyError("mindmap.create source must be an existing workspace .md or .txt file")
            if source_path.stat().st_size > 2 * 1024 * 1024:
                raise OmniBodyError("mindmap.create source exceeds 2 MiB")
            content = source_path.read_text(encoding="utf-8")
        inline_content_present = isinstance(content, str) and bool(content.strip())
        content_mode = "source_text" if source_path else "inline_text" if inline_content_present else "tree"
        if inline_content_present and not tree:
            # The first non-empty line is the map root.  Parse only the
            # remaining indented outline so the root is not duplicated as its
            # own child in Mermaid and OPML.
            outline_lines = [line for line in content.splitlines() if line.strip()]
            title = outline_lines[0].lstrip("#- *").strip()[:80] or title
            tree = self._parse_indented_tree("\n".join(outline_lines[1:]))
        lines = ["```mermaid", "mindmap", f"  root(({self._escape_mermaid(title)}))"]
        self._mindmap_lines(tree, lines, indent=4)
        lines.append("```")
        content = "\n".join(lines) + "\n"
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(content, encoding="utf-8")
        if opml_path is not None:
            opml = self._opml(title, tree)
            opml_path.write_text(opml, encoding="utf-8")
        mermaid_src = "\n".join(lines[1:-1])
        html_path.write_text(self._mindmap_html(title, mermaid_src), encoding="utf-8")
        return {
            "snapshots": snapshots,
            "output": self._file_evidence(output),
            "html": self._file_evidence(html_path),
            "opml": self._file_evidence(opml_path) if opml_path else None,
            "content_mode": content_mode,
            "source": str(source_path) if source_path else "",
            "requested_output": str(requested_output),
            "note": "MD 为 Mermaid 源码；HTML 可双击查看，首次渲染需联网加载 Mermaid 运行时。",
        }

    def _mindmap_html(self, title: Any, mermaid_src: str) -> str:
        safe_title = html_lib.escape(str(title or "Mindmap")[:120] or "Mindmap", quote=True)
        safe_source = html_lib.escape(str(mermaid_src or ""), quote=False)
        return (
            '<!DOCTYPE html>\n<html lang="zh-CN">\n<head>\n<meta charset="utf-8">\n'
            f"<title>{safe_title}</title>\n"
            '<meta name="viewport" content="width=device-width, initial-scale=1">\n'
            "<style>body{margin:0;font-family:'Microsoft YaHei',system-ui,sans-serif;background:#f7f8fa}"
            "header{padding:14px 20px;font-size:16px;font-weight:700;color:#1f2937;"
            "border-bottom:1px solid #e5e7eb;background:#fff}"
            "#map{padding:24px;overflow:auto}</style>\n</head>\n<body>\n"
            f"<header>{safe_title}</header>\n"
            f'<pre id="map" class="mermaid">\n{safe_source}\n</pre>\n'
            '<script type="module">\n'
            "const pre=document.getElementById('map');\n"
            "const source=pre?pre.textContent:'';\n"
            "try{\n"
            "  const mermaid=(await import('https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs')).default;\n"
            "  mermaid.initialize({startOnLoad:false,theme:'default',securityLevel:'strict'});\n"
            "  if(pre&&source){\n"
            "    pre.hidden=true;\n"
            "    const holder=document.createElement('div');\n"
            "    holder.className='mermaid';\n"
            "    holder.textContent=source;\n"
            "    pre.parentNode.insertBefore(holder,pre.nextSibling);\n"
            "    await mermaid.run({nodes:[holder]});\n"
            "  }\n"
            "}catch(error){\n"
            "  // MM-P2-GROUP-6: offline fallback — the <pre> keeps the raw\n"
            "  // mindmap text visible instead of leaving a blank page.\n"
            "  document.body.dataset.mermaidOffline='true';\n"
            "}\n"
            "</script>\n</body>\n</html>\n"
        )

    def _escape_mermaid(self, text: Any) -> str:
        value = re.sub(r"[\x00-\x1f\x7f]+", " ", str(text))
        value = re.sub(r"[()\[\]{}<>`\"]", "", value).replace(":::", " - ").replace("%%", "%")
        return re.sub(r"\s+", " ", value).strip() or "未命名"

    def _parse_indented_tree(self, text: str) -> list:
        """解析缩进文本为嵌套 dict 树"""
        lines = [ln for ln in text.split("\n") if ln.strip()]
        if not lines:
            return []
        root: list = []
        stack: list[tuple[int, Any]] = []  # (indent, parent)
        for raw_line in lines:
            line = raw_line.expandtabs(2)
            stripped = line.lstrip(" ")
            indent = len(line) - len(stripped)
            label = stripped.lstrip("- *#").strip()
            if not label:
                continue
            node = {label: []}
            while stack and stack[-1][0] >= indent:
                stack.pop()
            if stack:
                parent = stack[-1][1]
                if isinstance(parent, dict):
                    for v in parent.values():
                        if isinstance(v, list):
                            v.append(node)
                            break
                elif isinstance(parent, list):
                    parent.append(node)
            else:
                root.append(node)
            stack.append((indent, node))
        return root

    def _mindmap_lines(self, node: Any, lines: List[str], indent: int) -> None:
        pad = " " * indent
        if isinstance(node, dict):
            for k, v in node.items():
                lines.append(f"{pad}{self._escape_mermaid(k)}")
                self._mindmap_lines(v, lines, indent + 2)
        elif isinstance(node, list):
            for item in node:
                self._mindmap_lines(item, lines, indent)
        elif node is not None:
            lines.append(f"{pad}{self._escape_mermaid(node)}")

    def _opml(self, title: str, tree: Any) -> str:
        def esc(x: Any) -> str:
            value = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", str(x))
            return value.replace("&", "&amp;").replace('"', "&quot;").replace("<", "&lt;").replace(">", "&gt;")
        def emit(node: Any, level: int = 2) -> List[str]:
            pad = "  " * level
            out: List[str] = []
            if isinstance(node, dict):
                for k, v in node.items():
                    children = emit(v, level + 1)
                    if children:
                        out.append(f'{pad}<outline text="{esc(k)}">')
                        out.extend(children)
                        out.append(f'{pad}</outline>')
                    else:
                        out.append(f'{pad}<outline text="{esc(k)}"/>')
            elif isinstance(node, list):
                for item in node:
                    out.extend(emit(item, level))
            elif node is not None:
                out.append(f'{pad}<outline text="{esc(node)}"/>')
            return out
        lines = ['<?xml version="1.0" encoding="UTF-8"?>', '<opml version="2.0">', '<head>', f'<title>{esc(title)}</title>', '</head>', '<body>']
        lines.extend(emit(tree))
        lines.extend(['</body>', '</opml>'])
        return "\n".join(lines) + "\n"

    # ---------- pdf ----------

    def _action_pdf_extract_text(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        try:
            from pypdf import PdfReader
        except Exception as e:
            raise OmniBodyError("pypdf is required for pdf.extract_text") from e
        p = self._resolve(target, must_exist=True)
        reader = PdfReader(str(p))
        pages = []
        max_pages = int(args.get("max_pages", len(reader.pages)))
        for i, page in enumerate(reader.pages[:max_pages]):
            pages.append({"page": i + 1, "text": page.extract_text() or ""})
        return {"path": str(p), "page_count": len(reader.pages), "pages": pages}

    def _find_pdf_font(self, args: Dict[str, Any]) -> Optional[Path]:
        candidates: List[Path] = []
        explicit = str(args.get("font_path") or os.environ.get("TIANGONG_CJK_FONT") or "").strip()
        if explicit:
            candidates.append(Path(explicit).expanduser())
        # Prefer TrueType-outline fonts. Some Noto CJK TTC files contain CFF
        # outlines which ReportLab cannot embed even though the file exists.
        windows_root = str(os.environ.get("WINDIR") or os.environ.get("SystemRoot") or "").strip()
        if windows_root:
            font_root = Path(windows_root) / "Fonts"
            candidates.extend(font_root / name for name in ("simhei.ttf", "msyh.ttc", "simsun.ttc"))
        font_config = shutil.which("fc-match")
        if font_config:
            for family in ("Noto Sans CJK SC", "WenQuanYi Zen Hei", "AR PL UMing CN", "PingFang SC"):
                try:
                    probe = subprocess.run(
                        [font_config, "-f", "%{file}\n", family],
                        check=False,
                        capture_output=True,
                        text=True,
                        timeout=3,
                    )
                    candidate = probe.stdout.splitlines()[0].strip() if probe.stdout else ""
                    if candidate:
                        candidates.append(Path(candidate))
                except (OSError, subprocess.SubprocessError):
                    continue
        try:
            from reportlab.pdfbase.ttfonts import TTFont as _ProbeTTFont
        except Exception:
            _ProbeTTFont = None
        for path in candidates:
            if not (path.exists() and path.is_file()):
                continue
            resolved = path.resolve()
            if _ProbeTTFont is not None:
                try:
                    _ProbeTTFont("TiangongCJKProbe", str(resolved), subfontIndex=0)
                except TypeError:
                    try:
                        _ProbeTTFont("TiangongCJKProbe", str(resolved))
                    except Exception:
                        continue
                except Exception:
                    continue
            return resolved
        return None

    def _action_pdf_create_from_text(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        output = self._resolve(target or args.get("output") or "document.pdf")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        text = str(args.get("content") or args.get("text", ""))
        contains_cjk = bool(re.search(r"[\u3400-\u9fff]", text))
        try:
            from reportlab.lib.enums import TA_LEFT
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
            font_path = self._find_pdf_font(args) if contains_cjk else None
            font_name = "Helvetica"
            font_embedded = False
            if contains_cjk:
                if not font_path:
                    raise OmniBodyError("No embeddable CJK font was found. Set args.font_path or TIANGONG_CJK_FONT.")
                font_name = "TiangongCJK"
                try:
                    pdfmetrics.registerFont(TTFont(font_name, str(font_path), subfontIndex=0))
                except TypeError:
                    pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
                font_embedded = True
            styles = getSampleStyleSheet()
            body = ParagraphStyle("TiangongBody", parent=styles["BodyText"], fontName=font_name, fontSize=10.5, leading=16, wordWrap="CJK", alignment=TA_LEFT, spaceAfter=6)
            title_style = ParagraphStyle("TiangongTitle", parent=styles["Title"], fontName=font_name, fontSize=20, leading=28, wordWrap="CJK", spaceAfter=14)
            doc = SimpleDocTemplate(str(output), pagesize=A4, leftMargin=22*mm, rightMargin=22*mm, topMargin=20*mm, bottomMargin=20*mm, title=str(args.get("title") or ""))
            story = []
            title = str(args.get("title") or "").strip()
            if title:
                story.extend([Paragraph(title.replace("&", "&amp;"), title_style), Spacer(1, 6)])
            for line in text.splitlines() or [""]:
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") or "&nbsp;"
                story.append(Paragraph(safe, body))
            doc.build(story)
            writer = "reportlab_platypus"
        except ModuleNotFoundError:
            if contains_cjk:
                raise OmniBodyError("reportlab is required for CJK PDF generation")
            self._write_minimal_pdf(output, text)
            writer = "stdlib_minimal_pdf_latin1"
            font_path = None; font_name = "Helvetica"; font_embedded = False
        return {"snapshots": snapshots, "output": self._file_evidence(output), "writer": writer, "font_path": str(font_path or ""), "font_name": font_name, "font_embedded": font_embedded, "contains_cjk": contains_cjk}

    def _write_minimal_pdf(self, output: Path, text: str) -> None:
        # Small dependency-free fallback. It uses Helvetica Type1 and latin-1 replacement,
        # so it is for emergency v3 compatibility rather than high-quality CJK PDF layout.
        def pdf_escape(s: str) -> str:
            return s.encode("latin-1", "replace").decode("latin-1").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

        lines: List[str] = []
        for para in text.splitlines() or [""]:
            while len(para) > 88:
                lines.append(para[:88])
                para = para[88:]
            lines.append(para)
        page_height = 842
        page_width = 595
        per_page = 48
        pages = [lines[i:i + per_page] for i in range(0, len(lines), per_page)] or [[""]]
        objects: List[bytes] = []
        catalog_id = 1
        pages_id = 2
        font_id = 3
        page_ids = []
        content_ids = []
        # placeholders added below after page count is known
        objects.append(b"")
        objects.append(b"")
        objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        next_id = 4
        for page_lines in pages:
            page_id = next_id; content_id = next_id + 1; next_id += 2
            page_ids.append(page_id); content_ids.append(content_id)
            content = ["BT", "/F1 11 Tf", "50 790 Td", "14 TL"]
            for idx, line in enumerate(page_lines):
                if idx == 0:
                    content.append(f"({pdf_escape(line)}) Tj")
                else:
                    content.append(f"T* ({pdf_escape(line)}) Tj")
            content.append("ET")
            stream = "\n".join(content).encode("latin-1", "replace")
            objects.append(f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 {page_width} {page_height}] /Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>".encode("latin-1"))
            objects.append(f"<< /Length {len(stream)} >>\nstream\n".encode("latin-1") + stream + b"\nendstream")
        kids = " ".join(f"{pid} 0 R" for pid in page_ids)
        objects[0] = f"<< /Type /Catalog /Pages {pages_id} 0 R >>".encode("latin-1")
        objects[1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>".encode("latin-1")
        data = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets = [0]
        for obj_id, obj in enumerate(objects, 1):
            offsets.append(len(data))
            data.extend(f"{obj_id} 0 obj\n".encode("latin-1"))
            data.extend(obj)
            data.extend(b"\nendobj\n")
        xref_at = len(data)
        data.extend(f"xref\n0 {len(objects)+1}\n0000000000 65535 f \n".encode("latin-1"))
        for off in offsets[1:]:
            data.extend(f"{off:010d} 00000 n \n".encode("latin-1"))
        data.extend(f"trailer\n<< /Size {len(objects)+1} /Root {catalog_id} 0 R >>\nstartxref\n{xref_at}\n%%EOF\n".encode("latin-1"))
        output.write_bytes(bytes(data))

    # ---------- images ----------

    def _pil(self):
        try:
            from PIL import Image, ImageDraw, ImageFont
            return Image, ImageDraw, ImageFont
        except Exception as e:
            raise OmniBodyError("Pillow is required for image actions") from e

    def _action_image_info(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        Image, _, _ = self._pil()
        p = self._resolve(target, must_exist=True)
        with Image.open(p) as im:
            return {"path": str(p), "format": im.format, "mode": im.mode, "size": list(im.size), "width": im.width, "height": im.height}

    def _action_image_create_canvas(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        Image, _, _ = self._pil()
        output = self._resolve(target or args.get("output") or "canvas.png")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        w = int(args.get("width", 1080)); h = int(args.get("height", 1080))
        color = args.get("color", "white")
        im = Image.new(args.get("mode", "RGB"), (w, h), color=color)
        im.save(output)
        return {"snapshots": snapshots, "output": self._file_evidence(output), "size": [w, h]}

    def _action_image_resize(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        Image, _, _ = self._pil()
        src = self._resolve(target, must_exist=True)
        output = self._resolve(args.get("output") or src.with_name(src.stem + "_resized" + src.suffix).name)
        snapshots = self._snapshot(op_id, [output])
        with Image.open(src) as im:
            width = args.get("width"); height = args.get("height")
            if args.get("keep_ratio", True):
                im.thumbnail((int(width or im.width), int(height or im.height)))
                out = im.copy()
            else:
                out = im.resize((int(width or im.width), int(height or im.height)))
            output.parent.mkdir(parents=True, exist_ok=True)
            out.save(output)
        return {"snapshots": snapshots, "input": str(src), "output": self._file_evidence(output)}

    def _action_image_crop(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        Image, _, _ = self._pil()
        src = self._resolve(target, must_exist=True)
        output = self._resolve(args.get("output") or src.with_name(src.stem + "_crop" + src.suffix).name)
        box = args.get("box") or [args.get("left", 0), args.get("top", 0), args.get("right"), args.get("bottom")]
        if box[2] is None or box[3] is None:
            raise OmniBodyError("image.crop requires args.box=[left,top,right,bottom]")
        snapshots = self._snapshot(op_id, [output])
        with Image.open(src) as im:
            out = im.crop(tuple(map(int, box)))
            output.parent.mkdir(parents=True, exist_ok=True)
            out.save(output)
        return {"snapshots": snapshots, "output": self._file_evidence(output), "box": box}

    def _action_image_rotate(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        Image, _, _ = self._pil()
        src = self._resolve(target, must_exist=True)
        output = self._resolve(args.get("output") or src.with_name(src.stem + "_rotated" + src.suffix).name)
        snapshots = self._snapshot(op_id, [output])
        with Image.open(src) as im:
            out = im.rotate(float(args.get("degrees", 90)), expand=bool(args.get("expand", True)))
            output.parent.mkdir(parents=True, exist_ok=True)
            out.save(output)
        return {"snapshots": snapshots, "output": self._file_evidence(output)}

    def _action_image_add_text(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        Image, ImageDraw, _ = self._pil()
        src = self._resolve(target, must_exist=True)
        output = self._resolve(args.get("output") or src.with_name(src.stem + "_text" + src.suffix).name)
        snapshots = self._snapshot(op_id, [output])
        with Image.open(src).convert("RGBA") as im:
            draw = ImageDraw.Draw(im)
            font = _load_text_font(args.get("font"), int(args.get("font_size", 48)))
            if args.get("xy") is not None:
                xy = tuple(args["xy"])
            else:
                xy = (int(args.get("x", 40)), int(args.get("y", 40)))
            fill = args.get("fill") or args.get("color") or "black"
            draw.text(xy, str(args.get("text", "")), fill=fill, font=font)
            output.parent.mkdir(parents=True, exist_ok=True)
            im.convert("RGB").save(output)
        font_path = str(getattr(font, "path", "") or "")
        try:
            font_name = " / ".join(str(part) for part in font.getname())
        except Exception:
            font_name = type(font).__name__
        return {
            "snapshots": snapshots,
            "output": self._file_evidence(output),
            "font": {"path": font_path, "name": font_name, "scalable": bool(font_path)},
        }

    def _action_image_compose(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        Image, _, _ = self._pil()
        base = self._resolve(target, must_exist=True)
        overlay = self._resolve(args.get("overlay"), must_exist=True)
        output = self._resolve(args.get("output") or base.with_name(base.stem + "_composed" + base.suffix).name)
        snapshots = self._snapshot(op_id, [output])
        with Image.open(base).convert("RGBA") as im, Image.open(overlay).convert("RGBA") as ov:
            if args.get("overlay_size"):
                ov = ov.resize(tuple(map(int, args["overlay_size"])))
            if args.get("opacity") is not None:
                alpha = ov.getchannel("A")
                alpha = alpha.point(lambda v: int(v * float(args["opacity"])))
                ov.putalpha(alpha)
            im.alpha_composite(ov, dest=tuple(args.get("xy", [0, 0])))
            output.parent.mkdir(parents=True, exist_ok=True)
            im.convert("RGB").save(output)
        return {"snapshots": snapshots, "output": self._file_evidence(output)}

    def _action_image_convert(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        Image, _, _ = self._pil()
        src = self._resolve(target, must_exist=True)
        output = self._resolve(args.get("output") or src.with_suffix("." + args.get("format", "png").lower()).name)
        snapshots = self._snapshot(op_id, [output])
        with Image.open(src) as im:
            output.parent.mkdir(parents=True, exist_ok=True)
            im.save(output, format=args.get("format"))
        return {"snapshots": snapshots, "output": self._file_evidence(output)}

    # ---------- audio/video ----------

    def _action_audio_tone(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        if wave is None:
            raise OmniBodyError("audio.tone requires Python stdlib module wave; packaged runtime is missing wave.py")
        output = self._resolve(target or args.get("output") or "tone.wav")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        sample_rate = int(args.get("sample_rate", 44100))
        duration = float(args.get("duration", 1.0))
        freq = float(args.get("frequency", 440.0))
        volume = float(args.get("volume", 0.3))
        frames = int(sample_rate * duration)
        with wave.open(str(output), "w") as w:
            w.setnchannels(1); w.setsampwidth(2); w.setframerate(sample_rate)
            for i in range(frames):
                val = int(32767 * volume * math.sin(2 * math.pi * freq * i / sample_rate))
                w.writeframes(struct.pack("<h", val))
        return {"snapshots": snapshots, "output": self._file_evidence(output), "duration": duration, "frequency": freq}

    def _action_audio_trim(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        ffmpeg = self._require_ffmpeg()
        src = self._resolve(target, must_exist=True)
        output = self._resolve(args.get("output") or src.with_name(src.stem + "_trim" + src.suffix).name)
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        cmd = [ffmpeg, "-y", "-ss", str(args.get("start", 0)), "-i", str(src)]
        if args.get("duration") is not None:
            cmd += ["-t", str(args["duration"])]
        elif args.get("end") is not None:
            cmd += ["-to", str(args["end"])]
        cmd += [str(output)]
        res = self._run_subprocess(cmd, timeout=int(args.get("timeout", 120)))
        return {"snapshots": snapshots, "ffmpeg": res, "output": self._file_evidence(output), "success": res["ok"]}

    def _action_audio_concat(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        ffmpeg = self._require_ffmpeg()
        files = args.get("files") or ([] if target is None else [target])
        if len(files) < 2:
            raise OmniBodyError("audio.concat requires at least two files")
        srcs = [self._resolve(f, must_exist=True) for f in files]
        output = self._resolve(args.get("output") or "concat_audio.wav")
        snapshots = self._snapshot(op_id, [output])
        list_file = self.workspace / f".{op_id}_concat.txt"
        concat_lines = []
        for item in srcs:
            # ffmpeg concat demuxer quoting: wrap in single quotes and escape any single quote in the path.
            escaped = str(item).replace("'", "'\\''")
            concat_lines.append(f"file '{escaped}'\n")
        list_file.write_text("".join(concat_lines), encoding="utf-8")
        try:
            cmd = [ffmpeg, "-y", "-f", "concat", "-safe", "0", "-i", str(list_file), "-c", "copy", str(output)]
            res = self._run_subprocess(cmd, timeout=int(args.get("timeout", 120)))
        finally:
            try: list_file.unlink()
            except FileNotFoundError: pass
        return {"snapshots": snapshots, "ffmpeg": res, "output": self._file_evidence(output), "success": res["ok"]}

    def _action_video_info(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        src = self._resolve(target, must_exist=True)
        if self.ffprobe:
            cmd = [self.ffprobe, "-v", "error", "-show_format", "-show_streams", "-of", "json", str(src)]
            res = self._run_subprocess(cmd, timeout=int(args.get("timeout", 60)))
            info = {}
            if res["ok"]:
                try:
                    info = json.loads(res["stdout"])
                except Exception:
                    info = {"raw": res["stdout"]}
            return {"path": str(src), "ffprobe": res, "info": info, "success": res["ok"]}
        try:
            import imageio_ffmpeg  # type: ignore

            frames = imageio_ffmpeg.read_frames(str(src))
            try:
                info = dict(next(frames))
            finally:
                frames.close()
            return {
                "path": str(src),
                "probe_adapter": "imageio_ffmpeg",
                "info": info,
                "success": True,
            }
        except Exception as exc:
            raise OmniBodyError(
                "ffprobe not found and portable imageio-ffmpeg metadata probe failed"
            ) from exc

    def _action_video_cut(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        ffmpeg = self._require_ffmpeg()
        src = self._resolve(target, must_exist=True)
        output = self._resolve(args.get("output") or src.with_name(src.stem + "_cut" + src.suffix).name)
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        cmd = [ffmpeg, "-y", "-ss", str(args.get("start", 0)), "-i", str(src)]
        if args.get("duration") is not None:
            cmd += ["-t", str(args["duration"])]
        elif args.get("end") is not None:
            cmd += ["-to", str(args["end"])]
        if args.get("copy", True):
            cmd += ["-c", "copy"]
        cmd += [str(output)]
        res = self._run_subprocess(cmd, timeout=int(args.get("timeout", 240)))
        return {"snapshots": snapshots, "ffmpeg": res, "output": self._file_evidence(output), "success": res["ok"]}

    def _action_video_extract_audio(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        ffmpeg = self._require_ffmpeg()
        src = self._resolve(target, must_exist=True)
        output = self._resolve(args.get("output") or src.with_suffix(".wav").name)
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        cmd = [ffmpeg, "-y", "-i", str(src), "-vn", "-acodec", args.get("codec", "pcm_s16le"), str(output)]
        res = self._run_subprocess(cmd, timeout=int(args.get("timeout", 240)))
        return {"snapshots": snapshots, "ffmpeg": res, "output": self._file_evidence(output), "success": res["ok"]}

    def _action_video_add_audio(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        ffmpeg = self._require_ffmpeg()
        video = self._resolve(target, must_exist=True)
        audio = self._resolve(args.get("audio"), must_exist=True)
        output = self._resolve(args.get("output") or video.with_name(video.stem + "_audio" + video.suffix).name)
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        cmd = [ffmpeg, "-y", "-i", str(video), "-i", str(audio), "-map", "0:v:0", "-map", "1:a:0"]
        # ``-shortest`` can silently discard a sparse/single-frame video stream
        # when stream-copying slideshows. Preserve the full video by default;
        # callers may opt in when they explicitly want the shorter duration.
        if bool(args.get("shortest", False)):
            cmd.append("-shortest")
        cmd += ["-c:v", args.get("video_codec", "copy"), "-c:a", args.get("audio_codec", "aac"), str(output)]
        res = self._run_subprocess(cmd, timeout=int(args.get("timeout", 240)))
        evidence = self._file_evidence(output)
        stream_types: List[str] = []
        if res["ok"] and output.exists() and self.ffprobe:
            probe = self._run_subprocess(
                [self.ffprobe, "-v", "error", "-show_entries", "stream=codec_type", "-of", "json", str(output)],
                timeout=int(args.get("probe_timeout", 60)),
            )
            if probe["ok"]:
                try:
                    stream_types = [
                        str(item.get("codec_type") or "")
                        for item in (json.loads(probe["stdout"]).get("streams") or [])
                    ]
                except Exception:
                    stream_types = []
            if "video" not in stream_types or "audio" not in stream_types:
                res = dict(res)
                res["ok"] = False
                res["postcondition_error"] = "muxed output must contain both video and audio streams"
        return {
            "snapshots": snapshots,
            "ffmpeg": res,
            "output": evidence,
            "stream_types": stream_types,
            "success": bool(res["ok"]),
        }

    def _action_video_slideshow(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        ffmpeg = self._require_ffmpeg()
        images = args.get("images") or ([] if target is None else [target])
        if not images:
            raise OmniBodyError("video.slideshow requires args.images")
        srcs = [self._resolve(i, must_exist=True) for i in images]
        output = self._resolve(target or args.get("output") or "slideshow.mp4")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        # Create a temporary sequence of copied images with normalized names.
        # Operation IDs include timestamps/action names and can push the
        # sandboxed Windows path beyond MAX_PATH.  Keep temporary components
        # short while retaining collision resistance.
        tmp = self.workspace / f".slides_{hashlib.sha256(op_id.encode('utf-8')).hexdigest()[:12]}"
        tmp.mkdir(exist_ok=True)
        try:
            for i, p in enumerate(srcs):
                dst = tmp / f"slide_{i:04d}{p.suffix.lower()}"
                shutil.copy2(p, dst)
            first_suffix = srcs[0].suffix.lower()
            frame_rate = float(args.get("frame_rate", 1 / float(args.get("seconds_per_image", 2))))
            # The Windows ffmpeg builds used by portable/source runtimes do
            # not consistently implement ``-pattern_type glob``.  The
            # normalized sequence names are already deterministic.
            sequence = tmp / f"slide_%04d{first_suffix}"
            cmd = [ffmpeg, "-y", "-framerate", str(frame_rate), "-i", str(sequence), "-pix_fmt", "yuv420p", str(output)]
            res = self._run_subprocess(cmd, timeout=int(args.get("timeout", 240)))
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
        return {"snapshots": snapshots, "ffmpeg": res, "output": self._file_evidence(output), "success": res["ok"]}


    # ---------- portable app-bus adapters ----------

    def _safe_slug(self, value: Any, fallback: str = "item") -> str:
        raw = str(value or fallback).strip() or fallback
        raw = re.sub(r"[^a-zA-Z0-9_\-.\u4e00-\u9fff]+", "_", raw).strip("._")
        return raw[:80] or fallback

    class _HTMLTextExtractor(_HTMLParser if _HTMLParser is not None else object):
        def __init__(self) -> None:
            if _HTMLParser is not None:
                super().__init__()
            self.in_skip = False
            self.parts: List[str] = []
            self.title_parts: List[str] = []
            self.in_title = False
        def handle_starttag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
            tag = tag.lower()
            if tag in {"script", "style", "noscript"}:
                self.in_skip = True
            if tag == "title":
                self.in_title = True
            if tag in {"p", "br", "div", "section", "article", "li", "h1", "h2", "h3", "tr"}:
                self.parts.append("\n")
        def handle_endtag(self, tag: str) -> None:
            tag = tag.lower()
            if tag in {"script", "style", "noscript"}:
                self.in_skip = False
            if tag == "title":
                self.in_title = False
        def handle_data(self, data: str) -> None:
            if self.in_skip:
                return
            txt = re.sub(r"\s+", " ", data).strip()
            if not txt:
                return
            if self.in_title:
                self.title_parts.append(txt)
            self.parts.append(txt)
        def text(self) -> str:
            return re.sub(r"\n{3,}", "\n\n", " ".join(self.parts).replace(" \n ", "\n")).strip()
        def title(self) -> str:
            return " ".join(self.title_parts).strip()

    def _html_to_text(self, html_text: str) -> Tuple[str, str]:
        if _HTMLParser is None:
            title_match = re.search(r"(?is)<title[^>]*>(.*?)</title>", html_text or "")
            title = re.sub(r"\s+", " ", title_match.group(1)).strip() if title_match else ""
            cleaned = re.sub(r"(?is)<(script|style|noscript)[^>]*>.*?</\1>", " ", html_text or "")
            cleaned = re.sub(r"(?is)<br\s*/?>", "\n", cleaned)
            cleaned = re.sub(r"(?is)</p\s*>|</div\s*>|</li\s*>|</h[1-6]\s*>", "\n", cleaned)
            plain = re.sub(r"(?is)<[^>]+>", " ", cleaned)
            replacements = {"&nbsp;": " ", "&amp;": "&", "&lt;": "<", "&gt;": ">", "&quot;": '"', "&#39;": "'"}
            for src, dst in replacements.items():
                plain = plain.replace(src, dst)
            plain = re.sub(r"[ \t\r\f\v]+", " ", plain)
            plain = re.sub(r"\n\s*\n\s*\n+", "\n\n", plain)
            plain = "\n".join(line.strip() for line in plain.splitlines() if line.strip())
            return plain.strip(), title
        parser = self._HTMLTextExtractor()
        parser.feed(html_text)
        return parser.text(), parser.title()

    def _normalize_url(self, target: Optional[str]) -> str:
        url = str(target or "").strip()
        if not url:
            raise OmniBodyError("browser action requires target URL/path")
        if re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", url):
            return url
        # Existing local file path is treated as a file resource. Otherwise assume https URL.
        try:
            p = self._resolve(url, must_exist=False)
            if p.exists():
                return p.as_uri()
        except Exception:
            pass
        return "https://" + url

    def _browser_fetch(self, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        url = self._normalize_url(target)
        max_bytes = max(1, min(int(args.get("max_bytes", 5_000_000)), 20_000_000))
        timeout = max(1, min(int(args.get("timeout", 20)), 120))
        headers = dict(args.get("headers") or {})
        headers.setdefault(
            "User-Agent",
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/126.0.0.0 Safari/537.36",
        )
        headers.setdefault("Accept", "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8")
        headers.setdefault("Accept-Language", "zh-CN,zh;q=0.9,en;q=0.8")
        parsed = urllib.parse.urlparse(url)
        if parsed.scheme == "data":
            header, _, data = url.partition(",")
            raw = urllib.parse.unquote_to_bytes(data)
            if ";base64" in header:
                raw = base64.b64decode(data)
            raw = raw[:max_bytes]
            final_url = url[:96] + ("..." if len(url) > 96 else "")
            content_type = header.split(":", 1)[1].split(";", 1)[0] if ":" in header else "text/plain"
            status = 200
        elif parsed.scheme == "file":
            raw_path = urllib.request.url2pathname(parsed.path)
            p = self._resolve(raw_path, must_exist=True)
            raw = p.read_bytes()[:max_bytes]
            final_url = p.as_uri()
            content_type = mimetypes.guess_type(str(p))[0] or "application/octet-stream"
            status = 200
        elif parsed.scheme in {"http", "https"}:
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req, timeout=timeout) as resp:  # nosec - user-controlled browser fetch tool
                try:
                    from .network_capabilities import NetworkCapabilityError, read_bounded_http_body
                    raw = read_bounded_http_body(resp, max_bytes)
                except NetworkCapabilityError as exc:
                    raise OmniBodyError(str(exc)) from exc
                final_url = resp.geturl()
                if urllib.parse.urlparse(final_url).scheme not in {"http", "https"}:
                    raise OmniBodyError("browser redirect changed to a forbidden URL scheme")
                content_type = resp.headers.get("Content-Type", "application/octet-stream")
                status = getattr(resp, "status", 200)
        else:
            raise OmniBodyError(f"browser URL scheme is forbidden: {parsed.scheme or 'missing'}")
        declared_charset = re.search(r"(?i)\bcharset\s*=\s*[\"']?([^;\"'\s]+)", content_type)
        encoding = str(args.get("encoding") or (declared_charset.group(1) if declared_charset else "utf-8"))
        try:
            text = raw.decode(encoding, errors="replace")
        except LookupError:
            encoding = "utf-8"
            text = raw.decode(encoding, errors="replace")
        return {
            "url": url,
            "final_url": final_url,
            "status": status,
            "content_type": content_type,
            "encoding": encoding,
            "raw": raw,
            "text": text,
            "bytes": len(raw),
        }

    def _action_browser_chrome_goto(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        output_dir = self._resolve(args.get("output_dir") or "browser_snapshots")
        output_dir.mkdir(parents=True, exist_ok=True)
        fetched = self._browser_fetch(target, args)
        url_for_name = fetched.get("final_url") or fetched.get("url") or target or "page"
        slug = self._safe_slug(urllib.parse.urlparse(str(url_for_name)).netloc + "_" + urllib.parse.urlparse(str(url_for_name)).path.replace("/", "_"), "page")
        ext = ".html" if "html" in str(fetched.get("content_type", "")).lower() else ".txt"
        html_path = output_dir / f"{op_id}_{slug}{ext}"
        text_path = output_dir / f"{op_id}_{slug}.txt"
        snapshots = self._snapshot(op_id, [html_path, text_path])
        html_path.write_bytes(fetched["raw"])
        text, title = self._html_to_text(fetched["text"])
        text_path.write_text(text or fetched["text"][:200000], encoding="utf-8")
        body_probe = (text or fetched.get("text") or "")[:3000].casefold()
        wall_markers = (
            "安全验证",
            "访问验证",
            "人机验证",
            "captcha",
            "verify you are",
            "checking your browser",
            "请完成验证",
            "滑动验证",
            "访问过于频繁",
            "环境异常",
            "安全检测",
        )
        bot_wall = any(marker.casefold() in body_probe for marker in wall_markers)
        result = {
            "snapshots": snapshots,
            "url": fetched["url"],
            "final_url": fetched["final_url"],
            "status": fetched["status"],
            "content_type": fetched["content_type"],
            "title": title,
            "html_snapshot": self._file_evidence(html_path),
            "text_snapshot": self._file_evidence(text_path),
            "text_preview": (text or fetched["text"])[: int(args.get("preview_chars", 1200))],
            "evidence": self._file_evidence(html_path),
        }
        if bot_wall:
            result["bot_wall_suspected"] = True
            result["note"] = "站点疑似返回反爬或安全验证页；请改读其他来源，不要重复抓取同一地址。"
        return result

    def _action_browser_chrome_open(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        return self._action_browser_chrome_goto(op_id, target, args)

    def _action_browser_open(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        return self._action_browser_chrome_goto(op_id, target, args)

    def _action_http_get(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        url = str(args.get("url") or target or "").strip()
        if not url:
            raise OmniBodyError("http.get requires target or args.url")
        fetched = self._browser_fetch(url, args)
        text = fetched.get("text", "")
        max_chars = int(args.get("max_chars", 200000))
        return {
            "url": fetched["url"],
            "final_url": fetched["final_url"],
            "status": fetched["status"],
            "content_type": fetched["content_type"],
            "body": text[:max_chars],
            "text": text[:max_chars],
            "bytes": fetched["bytes"],
            "evidence": {
                "url": fetched["final_url"],
                "status": fetched["status"],
                "content_type": fetched["content_type"],
                "bytes": fetched["bytes"],
            },
        }

    def _action_browser_chrome_extract_text(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        # If target is a local file, parse it; otherwise fetch first.
        p: Optional[Path] = None
        try:
            p = self._resolve(target, must_exist=False) if target else None
        except Exception:
            p = None
        if p and p.exists():
            html_text = p.read_text(encoding=args.get("encoding", "utf-8"), errors="replace")
            source = self._file_evidence(p)
        else:
            fetched = self._browser_fetch(target, args)
            html_text = fetched["text"]
            source = {"url": fetched["final_url"], "status": fetched["status"], "bytes": fetched["bytes"]}
        text, title = self._html_to_text(html_text)
        return {"title": title, "text": text[: int(args.get("max_chars", 200000))], "source": source, "evidence": source}

    def _action_browser_chrome_extract_dom(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        fetched = self._browser_fetch(target, args)
        return {"url": fetched["final_url"], "status": fetched["status"], "html": fetched["text"][: int(args.get("max_chars", 200000))], "content_type": fetched["content_type"]}

    def _action_browser_chrome_download(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        fetched = self._browser_fetch(target, args)
        filename = args.get("filename")
        if not filename:
            path_name = Path(urllib.parse.urlparse(str(fetched.get("final_url") or fetched.get("url"))).path).name
            filename = path_name or "download.bin"
        output = self._resolve(args.get("output") or Path("downloads") / filename)
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        fd, temp_name = tempfile.mkstemp(prefix=f".{output.name}.", suffix=".download", dir=str(output.parent))
        os.close(fd)
        temp_output = Path(temp_name)
        try:
            temp_output.write_bytes(fetched["raw"])
            os.replace(temp_output, output)
        finally:
            temp_output.unlink(missing_ok=True)
        return {"snapshots": snapshots, "url": fetched["final_url"], "status": fetched["status"], "output": self._file_evidence(output), "evidence": self._file_evidence(output)}

    def _action_web_download(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        return self._action_browser_chrome_download(op_id, target or args.get("url"), args)

    def _action_browser_chrome_pdf_print(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        extracted = self._action_browser_chrome_extract_text(op_id, target, args)
        output = args.get("output") or "browser_page.pdf"
        return self._action_pdf_create_from_text(op_id, output, {"text": extracted.get("text", ""), "title": extracted.get("title", "Browser page")})

    def _action_browser_chrome_screenshot(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        Image, ImageDraw, _ = self._pil()
        extracted = self._action_browser_chrome_extract_text(op_id, target, args)
        # MM-FE-05: honor the target path contract like every other action;
        # args.output is only a fallback for callers that pass it explicitly.
        output = self._resolve(target or args.get("output") or "browser_snapshot.png")
        snapshots = self._snapshot(op_id, [output])
        w, h = int(args.get("width", 1280)), int(args.get("height", 1600))
        im = Image.new("RGB", (w, h), color=args.get("background", "white"))
        draw = ImageDraw.Draw(im)
        y = 30
        for line in (str(extracted.get("title") or "Browser snapshot") + "\n\n" + str(extracted.get("text") or "")).splitlines():
            for chunk in [line[i:i+95] for i in range(0, len(line), 95)] or [""]:
                if y > h - 40:
                    break
                draw.text((30, y), chunk, fill=args.get("fill", "black"))
                y += 22
            if y > h - 40:
                break
        output.parent.mkdir(parents=True, exist_ok=True)
        im.save(output)
        return {"snapshots": snapshots, "output": self._file_evidence(output), "source": extracted.get("source"), "evidence": self._file_evidence(output), "note": "Portable text-image screenshot; use browser_driver for real viewport rendering."}

    def _extract_search_results(self, html_text: str) -> List[Dict[str, str]]:
        """Parse common search-result pages into stable title/url/snippet rows."""

        source = str(html_text or "")
        results: List[Dict[str, str]] = []
        seen: set[str] = set()

        def clean(value: str) -> str:
            without_tags = re.sub(r"(?is)<[^>]+>", " ", value or "")
            return re.sub(r"\s+", " ", html_lib.unescape(without_tags)).strip()

        def push(url: str, title: str, snippet: str = "") -> None:
            decoded_url = html_lib.unescape(str(url or "")).strip()
            clean_title = clean(title)
            parsed = urllib.parse.urlparse(decoded_url)
            if parsed.scheme not in {"http", "https"} or not parsed.netloc or not clean_title:
                return
            host = (parsed.hostname or "").casefold()
            if host in {
                "bing.com",
                "www.bing.com",
                "cn.bing.com",
                "google.com",
                "www.google.com",
                "duckduckgo.com",
                "www.duckduckgo.com",
            }:
                return
            key = urllib.parse.urlunsplit((parsed.scheme, parsed.netloc, parsed.path, parsed.query, ""))
            if key in seen:
                return
            seen.add(key)
            results.append(
                {
                    "title": clean_title[:200],
                    "url": decoded_url,
                    "snippet": clean(snippet)[:500],
                }
            )

        for block in re.findall(r'(?is)<li[^>]*class=["\'][^"\']*\bb_algo\b[^"\']*["\'][^>]*>.*?</li>', source):
            link = re.search(
                r'(?is)<h2[^>]*>\s*<a[^>]*href=["\']([^"\']+)["\'][^>]*>(.*?)</a>',
                block,
            )
            if not link:
                continue
            snippet = re.search(r"(?is)<p[^>]*>(.*?)</p>", block)
            push(link.group(1), link.group(2), snippet.group(1) if snippet else "")
        if not results:
            for href, label in re.findall(
                r'(?is)<a[^>]*href=["\'](https?://[^"\']+)["\'][^>]*>(.*?)</a>',
                source,
            ):
                if len(clean(label)) >= 8:
                    push(href, label)
        return results[:10]

    def _action_browser_search_web(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        query = str(args.get("query") or target or "").strip()
        if not query:
            raise OmniBodyError("browser.search_web requires query")
        endpoint = str(args.get("search_url") or os.environ.get("OMNI_SEARCH_URL") or "https://www.bing.com/search?q={query}")
        url = endpoint.replace("{query}", urllib.parse.quote(query))
        fetched = self._browser_fetch(url, args)
        results = self._extract_search_results(fetched.get("text") or "")
        evidence = {
            "url": fetched["final_url"],
            "status": fetched["status"],
            "bytes": fetched["bytes"],
        }
        if results:
            return {
                "query": query,
                "url": fetched["url"],
                "final_url": fetched["final_url"],
                "status": fetched["status"],
                "results": results,
                "result_count": len(results),
                "note": "结构化搜索结果包含 title、url、snippet；可用 web.read 读取具体网址。",
                "evidence": evidence,
            }
        text, title = self._html_to_text(fetched.get("text") or "")
        return {
            "query": query,
            "url": fetched["url"],
            "final_url": fetched["final_url"],
            "status": fetched["status"],
            "results": [],
            "result_count": 0,
            "title": title,
            "text_preview": (text or fetched.get("text") or "")[: int(args.get("preview_chars", 2000))],
            "note": "未解析到结构化结果，已返回页面文本预览。",
            "evidence": evidence,
        }

    def _design_paths(self, target: Optional[str]) -> Tuple[Path, Path]:
        raw = target or "design.omni_ps.json"
        p = self._resolve(raw)
        if p.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}:
            project = p.with_suffix(p.suffix + ".omni_ps.json")
            preview = p
        else:
            project = p if p.suffix.lower().endswith(".json") else p.with_suffix(p.suffix + ".omni_ps.json")
            preview = project.with_suffix(".png")
        return project, preview

    def _load_design_project(self, project: Path, args: Dict[str, Any]) -> Dict[str, Any]:
        if project.exists():
            return json.loads(project.read_text(encoding="utf-8"))
        return {"schema": "tiangong.omni.portable_photoshop.v1", "width": int(args.get("width", 1080)), "height": int(args.get("height", 1080)), "background": args.get("background", "white"), "layers": []}

    def _render_design_project(self, project_data: Dict[str, Any], preview: Path) -> Dict[str, Any]:
        Image, ImageDraw, _ = self._pil()
        w, h = int(project_data.get("width", 1080)), int(project_data.get("height", 1080))
        im = Image.new("RGBA", (w, h), color=project_data.get("background", "white"))
        draw = ImageDraw.Draw(im)
        for layer in project_data.get("layers", []):
            if not layer.get("visible", True):
                continue
            kind = layer.get("kind", "blank")
            x, y = int(layer.get("x", 0)), int(layer.get("y", 0))
            if kind in {"text", "title"}:
                draw.text((x, y), str(layer.get("text", layer.get("name", ""))), fill=layer.get("fill", "black"))
            elif kind in {"rect", "rectangle", "shape"}:
                ww, hh = int(layer.get("width", 300)), int(layer.get("height", 120))
                draw.rectangle([x, y, x+ww, y+hh], fill=layer.get("fill", "#D9EAF7"), outline=layer.get("outline", None))
            elif kind == "image" and layer.get("path"):
                src = self._resolve(layer.get("path"), must_exist=True)
                with Image.open(src).convert("RGBA") as overlay:
                    if layer.get("width") or layer.get("height"):
                        overlay.thumbnail((int(layer.get("width", overlay.width)), int(layer.get("height", overlay.height))))
                    im.alpha_composite(overlay, (x, y))
        preview.parent.mkdir(parents=True, exist_ok=True)
        im.convert("RGB").save(preview)
        return self._file_evidence(preview)

    def _save_design_project(self, op_id: str, project: Path, preview: Path, data: Dict[str, Any]) -> Dict[str, Any]:
        snapshots = self._snapshot(op_id, [project, preview])
        project.parent.mkdir(parents=True, exist_ok=True)
        data["updated_at"] = time.time()
        project.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        preview_ev = self._render_design_project(data, preview)
        return {"snapshots": snapshots, "project": self._file_evidence(project), "preview": preview_ev, "evidence": preview_ev, "portable_fallback": True}

    def _action_adobe_photoshop_document_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        project, preview = self._design_paths(target)
        data = {"schema": "tiangong.omni.portable_photoshop.v1", "width": int(args.get("width", 1080)), "height": int(args.get("height", 1080)), "background": args.get("background", "white"), "layers": []}
        if args.get("title"):
            data["layers"].append({"id": "title", "name": "title", "kind": "text", "text": str(args.get("title")), "x": 60, "y": 60, "fill": args.get("fill", "black")})
        return self._save_design_project(op_id, project, preview, data)

    def _action_adobe_photoshop_document_open(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        project, preview = self._design_paths(target)
        if not project.exists():
            raise FileNotFoundError(str(project))
        data = json.loads(project.read_text(encoding="utf-8"))
        return {"project": self._file_evidence(project), "preview": self._file_evidence(preview), "layer_count": len(data.get("layers", [])), "layers": data.get("layers", [])[:50], "evidence": self._file_evidence(project)}

    def _action_adobe_photoshop_layer_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        project, preview = self._design_paths(target)
        data = self._load_design_project(project, args)
        layer = dict(args)
        layer.setdefault("id", self._safe_slug(layer.get("name") or f"layer_{len(data.get('layers', []))+1}"))
        layer.setdefault("name", layer["id"])
        layer.setdefault("kind", "text" if layer.get("text") else "rectangle" if layer.get("fill") else "blank")
        layer.setdefault("x", 40 + 20 * len(data.get("layers", [])))
        layer.setdefault("y", 40 + 28 * len(data.get("layers", [])))
        layer.setdefault("visible", True)
        data.setdefault("layers", []).append(layer)
        return self._save_design_project(op_id, project, preview, data)

    def _action_adobe_photoshop_text_add(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        args = dict(args)
        args["kind"] = "text"
        args.setdefault("text", args.get("content") or args.get("name") or "Text")
        return self._action_adobe_photoshop_layer_create(op_id, target, args)

    def _action_adobe_photoshop_layer_update(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        project, preview = self._design_paths(target)
        data = self._load_design_project(project, args)
        layer_id = str(args.get("id") or args.get("name") or "").strip()
        if not layer_id:
            raise OmniBodyError("layer.update requires args.id or args.name")
        updated = False
        for layer in data.get("layers", []):
            if layer.get("id") == layer_id or layer.get("name") == layer_id:
                for k, v in args.items():
                    if k not in {"id"}:
                        layer[k] = v
                updated = True
                break
        if not updated:
            raise OmniBodyError(f"layer not found: {layer_id}")
        return self._save_design_project(op_id, project, preview, data)

    def _action_adobe_photoshop_export_png(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        project, preview = self._design_paths(target)
        if not preview.exists():
            data = self._load_design_project(project, args)
            self._render_design_project(data, preview)
        output = self._resolve(args.get("output") or preview.name)
        snapshots = self._snapshot(op_id, [output])
        self._copy_path(preview, output, overwrite=True)
        return {"snapshots": snapshots, "output": self._file_evidence(output), "project": self._file_evidence(project), "portable_fallback": True, "evidence": self._file_evidence(output)}

    def _action_adobe_photoshop_image_resize(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        return self._action_image_resize(op_id, target, args)

    def _action_adobe_photoshop_image_crop(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        return self._action_image_crop(op_id, target, args)

    def _jianying_project_path(self, target: Optional[str], args: Dict[str, Any]) -> Path:
        raw = args.get("project") or target or "jianying_project.omni_jy.json"
        p = self._resolve(raw)
        if p.suffix.lower() != ".json":
            p = p.with_suffix(p.suffix + ".omni_jy.json")
        return p

    def _load_jianying_project(self, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._jianying_project_path(target, args)
        if p.exists():
            return json.loads(p.read_text(encoding="utf-8"))
        return {"schema": "tiangong.omni.portable_jianying.v1", "title": str(args.get("title") or target or "Jianying Project"), "media": [], "cuts": [], "subtitles": [], "music": None, "cover": None}

    def _save_jianying_project(self, op_id: str, path: Path, data: Dict[str, Any]) -> Dict[str, Any]:
        snapshots = self._snapshot(op_id, [path])
        path.parent.mkdir(parents=True, exist_ok=True)
        data["updated_at"] = time.time()
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return {"snapshots": snapshots, "project": self._file_evidence(path), "evidence": self._file_evidence(path), "portable_fallback": True}

    def _action_jianying_project_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._jianying_project_path(target, args)
        data = {"schema": "tiangong.omni.portable_jianying.v1", "title": str(args.get("title") or target or "Jianying Project"), "media": [], "cuts": [], "subtitles": [], "music": None, "cover": None, "canvas": args.get("canvas", "1080x1920")}
        return self._save_jianying_project(op_id, p, data)

    def _action_jianying_media_import(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._jianying_project_path(target, args)
        data = self._load_jianying_project(str(p), args)
        media = args.get("media") or args.get("input") or args.get("path")
        if not media:
            raise OmniBodyError("jianying.media.import requires args.media/input/path")
        self._resolve(str(media), must_exist=True)
        data.setdefault("media", []).append(str(media))
        return self._save_jianying_project(op_id, p, data)

    def _action_jianying_timeline_cut(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._jianying_project_path(target, args)
        data = self._load_jianying_project(str(p), args)
        data.setdefault("cuts", []).append({"input": args.get("input"), "start": float(args.get("start", 0)), "end": args.get("end"), "duration": args.get("duration")})
        return self._save_jianying_project(op_id, p, data)

    def _action_jianying_subtitle_add(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._jianying_project_path(target, args)
        data = self._load_jianying_project(str(p), args)
        data.setdefault("subtitles", []).append({"text": str(args.get("text") or args.get("content") or ""), "start": float(args.get("start", 0)), "end": args.get("end")})
        return self._save_jianying_project(op_id, p, data)

    def _action_jianying_music_add(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._jianying_project_path(target, args)
        data = self._load_jianying_project(str(p), args)
        audio = args.get("audio") or args.get("music") or args.get("path")
        if audio:
            self._resolve(str(audio), must_exist=True)
        data["music"] = audio
        return self._save_jianying_project(op_id, p, data)

    def _action_jianying_cover_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        output = args.get("output") or "jianying_cover.png"
        canvas = self._action_image_create_canvas(op_id, output, {"width": int(args.get("width", 1080)), "height": int(args.get("height", 1920)), "color": args.get("background", "#111111")})
        self._action_image_add_text(op_id, output, {"text": str(args.get("title") or target or "视频封面"), "x": int(args.get("x", 80)), "y": int(args.get("y", 160)), "fill": args.get("fill", "white"), "output": output})
        return {"output": self._file_evidence(self._resolve(output, must_exist=True)), "canvas": canvas, "portable_fallback": True, "evidence": self._file_evidence(self._resolve(output, must_exist=True))}

    def _action_jianying_export_mp4(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        ffmpeg = self._require_ffmpeg()
        output = self._resolve(args.get("output") or "jianying_export.mp4")
        snapshots = self._snapshot(op_id, [output])
        project = self._load_jianying_project(target, args)
        input_video = args.get("input") or (project.get("media") or [None])[0]
        images = args.get("images") or [m for m in project.get("media", []) if str(m).lower().endswith((".png", ".jpg", ".jpeg", ".webp"))]
        audio = args.get("audio") or project.get("music")
        if input_video and str(input_video).lower().endswith((".mp4", ".mov", ".mkv", ".avi", ".webm")):
            src = self._resolve(str(input_video), must_exist=True)
            cut_args = {"output": str(output), "start": args.get("start", 0), "duration": args.get("duration"), "end": args.get("end"), "copy": bool(args.get("copy", True)), "timeout": args.get("timeout", 240)}
            if project.get("cuts") and not args.get("duration") and not args.get("end"):
                first = project["cuts"][0]
                cut_args.update({"start": first.get("start", 0), "duration": first.get("duration"), "end": first.get("end")})
            res = self._action_video_cut(op_id, str(src), cut_args)
        else:
            if not images:
                # Create a title-card fallback; this is an actual MP4, not a false success.
                cover = self._resolve(args.get("cover") or project.get("cover") or "jianying_title_card.png")
                self._action_image_create_canvas(op_id, str(cover), {"width": int(args.get("width", 1080)), "height": int(args.get("height", 1920)), "color": args.get("background", "#111111")})
                title = args.get("title") or project.get("title") or str(target or "Jianying Export")
                subtitle = "\n".join([s.get("text", "") for s in project.get("subtitles", []) if s.get("text")]) or str(args.get("subtitle") or "")
                self._action_image_add_text(op_id, str(cover), {"text": str(title) + ("\n" + subtitle if subtitle else ""), "x": 80, "y": 160, "fill": args.get("fill", "white"), "output": str(cover)})
                images = [str(cover)]
            res = self._action_video_slideshow(op_id, None, {"images": images, "output": str(output), "seconds_per_image": float(args.get("seconds_per_image", args.get("duration", 3))), "timeout": args.get("timeout", 240)})
        final_output = output
        mux_result = None
        if audio and output.exists():
            audio_output = output.with_name(output.stem + "_with_audio" + output.suffix)
            mux_result = self._action_video_add_audio(op_id, str(output), {"audio": str(audio), "output": str(audio_output), "timeout": args.get("timeout", 240)})
            if mux_result.get("success"):
                final_output = audio_output
        info = self._action_video_info(op_id, str(final_output), {}) if final_output.exists() else {}
        return {"snapshots": snapshots, "render": res, "audio_mux": mux_result, "output": self._file_evidence(final_output), "video_info": info.get("info"), "portable_fallback": True, "success": bool(final_output.exists()), "evidence": self._file_evidence(final_output)}

    def _feishu_local_path(self, target: Optional[str], args: Dict[str, Any], suffix: str = ".md") -> Path:
        title = args.get("title") or target or "feishu_doc"
        if target and str(target).strip():
            raw = str(target)
            p = self._resolve(raw)
            if p.suffix:
                return p.with_suffix(suffix)
        return self._resolve(Path("feishu_docs") / (self._safe_slug(title, "feishu_doc") + suffix))

    def _action_feishu_docs_doc_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        # A real Feishu API adapter should be mounted by the host with credentials.
        # Portable fallback creates local markdown + docx as a truthful deliverable.
        title = str(args.get("title") or target or "飞书文档")
        body = args.get("content") or args.get("body") or ""
        if isinstance(body, list):
            body = "\n\n".join(map(str, body))
        md = self._feishu_local_path(target, {**args, "title": title}, ".md")
        docx = md.with_suffix(".docx")
        snapshots = self._snapshot(op_id, [md, docx])
        md.parent.mkdir(parents=True, exist_ok=True)
        md.write_text(f"# {title}\n\n{body}\n", encoding="utf-8")
        self._action_docx_create(op_id, str(docx), {"title": title, "sections": [{"heading": "正文", "paragraphs": str(body).splitlines() or [""]}]})
        return {"snapshots": snapshots, "remote_created": False, "mode": "local_feishu_fallback", "markdown": self._file_evidence(md), "docx": self._file_evidence(docx), "evidence": self._file_evidence(docx), "warning": "No Feishu credentials/API adapter in portable core; created local Feishu-compatible deliverables."}

    def _action_feishu_docs_doc_read(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        return self._action_file_read(op_id, target, args)

    def _action_feishu_docs_doc_update(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._feishu_local_path(target, args, ".md")
        existing = p.read_text(encoding="utf-8") if p.exists() else ""
        content = str(args.get("content") or args.get("body") or "")
        mode = args.get("mode", "append")
        new_content = content if mode == "replace" else existing + ("\n\n" if existing and content else "") + content
        return self._action_file_write(op_id, str(p), {"content": new_content})

    def _action_feishu_docs_export_docx(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target, must_exist=True)
        text = p.read_text(encoding=args.get("encoding", "utf-8"), errors="replace")
        title = args.get("title") or p.stem
        output = args.get("output") or p.with_suffix(".docx").name
        return self._action_docx_create(op_id, output, {"title": title, "sections": [{"heading": "正文", "paragraphs": text.splitlines()}]})

    def _action_feishu_docs_export_pdf(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        p = self._resolve(target, must_exist=True)
        text = p.read_text(encoding=args.get("encoding", "utf-8"), errors="replace")
        output = args.get("output") or p.with_suffix(".pdf").name
        return self._action_pdf_create_from_text(op_id, output, {"title": args.get("title") or p.stem, "text": text})

    def _action_audio_tts(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        text = str(args.get("text") or args.get("content") or target or "").strip()
        if not text:
            raise OmniBodyError("audio.tts requires args.text")
        output = self._resolve(args.get("output") or "tts.wav")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        espeak = shutil.which("espeak") or shutil.which("espeak-ng")
        if espeak:
            voice = str(args.get("voice") or args.get("language") or "en")
            speed = str(args.get("speed") or args.get("words_per_minute") or "150")
            cmd = [espeak, "-v", voice, "-s", speed, "-w", str(output), text]
            res = self._run_subprocess(cmd, timeout=int(args.get("timeout", 120)))
            return {"snapshots": snapshots, "engine": Path(espeak).name, "subprocess": res, "output": self._file_evidence(output), "success": res["ok"] and output.exists(), "evidence": self._file_evidence(output)}
        if sys.platform == "darwin" and shutil.which("say"):
            aiff = output.with_suffix(".aiff")
            res = self._run_subprocess(["say", "-o", str(aiff), text], timeout=int(args.get("timeout", 120)))
            if res["ok"] and aiff.exists():
                if self.ffmpeg:
                    conv = self._run_subprocess([self.ffmpeg, "-y", "-i", str(aiff), str(output)], timeout=120)
                    try: aiff.unlink()
                    except Exception: pass
                    return {"snapshots": snapshots, "engine": "say+ffmpeg", "subprocess": conv, "output": self._file_evidence(output), "success": conv["ok"], "evidence": self._file_evidence(output)}
                aiff.rename(output)
                return {"snapshots": snapshots, "engine": "say", "output": self._file_evidence(output), "success": True, "evidence": self._file_evidence(output)}
        if os.name == "nt":
            powershell = shutil.which("powershell") or shutil.which("pwsh")
            if powershell:
                res = _run_windows_sapi_broker(
                    powershell,
                    text=text,
                    output=output,
                    voice=str(args.get("voice") or args.get("language") or ""),
                    rate=int(args.get("rate", 0)),
                    timeout=int(args.get("timeout", 120)),
                )
                wave_valid = False
                if res["ok"] and output.is_file() and output.stat().st_size > 44:
                    header = output.read_bytes()[:12]
                    wave_valid = header[:4] == b"RIFF" and header[8:12] == b"WAVE"
                if wave_valid:
                    evidence = self._file_evidence(output)
                    return {
                        "snapshots": snapshots,
                        "engine": "windows-sapi",
                        "subprocess": res,
                        "output": evidence,
                        "success": True,
                        "evidence": evidence,
                    }
                raise OmniBodyError(
                    "Windows SAPI did not create a valid WAV file"
                    + (f": {res.get('stderr', '')[-500:]}" if res.get("stderr") else "")
                )
        raise OmniBodyError("No local TTS engine found. Install espeak/espeak-ng or mount cloud TTS adapter.")

    def _action_elevenlabs_tts_create(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        args = dict(args)
        args.setdefault("text", args.get("text") or target)
        result = self._action_audio_tts(op_id, target, args)
        result["provider"] = "portable_local_tts_fallback_not_elevenlabs"
        return result

    def _desktop_enabled(self, args: Dict[str, Any]) -> bool:
        return str(os.environ.get("OMNI_DESKTOP_ENABLE", "")).lower() in {"1", "true", "yes", "on"} or bool(args.get("enable_desktop"))

    def _pyautogui(self, args: Dict[str, Any]):
        if not self._desktop_enabled(args):
            raise OmniBodyError("Desktop actuation disabled. Set OMNI_DESKTOP_ENABLE=1 or args.enable_desktop=true after user approval.")
        try:
            import pyautogui  # type: ignore
            return pyautogui
        except Exception as e:
            raise OmniBodyError("pyautogui is required for desktop click/type/hotkey") from e

    def _action_desktop_screenshot(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        output = self._resolve(target or args.get("output") or "desktop_screenshot.png")
        snapshots = self._snapshot(op_id, [output])
        output.parent.mkdir(parents=True, exist_ok=True)
        try:
            from PIL import ImageGrab  # type: ignore
            im = ImageGrab.grab()
            im.save(output)
            return {"snapshots": snapshots, "engine": "PIL.ImageGrab", "output": self._file_evidence(output), "evidence": self._file_evidence(output)}
        except Exception:
            if not self._desktop_enabled(args):
                raise OmniBodyError("Desktop screenshot unavailable in this environment. Provide display access or enable desktop adapter.")
            try:
                import pyautogui  # type: ignore
                im = pyautogui.screenshot()
                im.save(output)
                return {"snapshots": snapshots, "engine": "pyautogui", "output": self._file_evidence(output), "evidence": self._file_evidence(output)}
            except Exception as e:
                raise OmniBodyError(f"Desktop screenshot failed: {e}") from e

    def _action_desktop_click(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        pyautogui = self._pyautogui(args)
        x = int(args.get("x", 0)); y = int(args.get("y", 0))
        pyautogui.click(x=x, y=y, clicks=int(args.get("clicks", 1)), interval=float(args.get("interval", 0.0)), button=str(args.get("button", "left")))
        return {"clicked": {"x": x, "y": y}, "engine": "pyautogui", "evidence": {"x": x, "y": y}}

    def _action_desktop_type(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        pyautogui = self._pyautogui(args)
        text = str(args.get("text") or target or "")
        pyautogui.write(text, interval=float(args.get("interval", 0.0)))
        return {"typed_chars": len(text), "engine": "pyautogui", "evidence": {"typed_chars": len(text)}}

    def _action_desktop_hotkey(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        pyautogui = self._pyautogui(args)
        keys = args.get("keys") or ([target] if target else [])
        if isinstance(keys, str):
            keys = [k.strip() for k in keys.split("+") if k.strip()]
        if not keys:
            raise OmniBodyError("desktop.hotkey requires args.keys")
        pyautogui.hotkey(*keys)
        return {"keys": keys, "engine": "pyautogui", "evidence": {"keys": keys}}

    def _action_windows_desktop_screenshot(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_screenshot(op_id, target, args)
    def _action_windows_desktop_click(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_click(op_id, target, args)
    def _action_windows_desktop_type(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_type(op_id, target, args)
    def _action_windows_desktop_hotkey(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_hotkey(op_id, target, args)
    def _action_macos_desktop_screenshot(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_screenshot(op_id, target, args)
    def _action_macos_desktop_click(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_click(op_id, target, args)
    def _action_macos_desktop_type(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_type(op_id, target, args)
    def _action_macos_desktop_hotkey(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_hotkey(op_id, target, args)
    def _action_linux_desktop_screenshot(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_screenshot(op_id, target, args)
    def _action_linux_desktop_click(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_click(op_id, target, args)
    def _action_linux_desktop_type(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_type(op_id, target, args)
    def _action_linux_desktop_hotkey(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]: return self._action_desktop_hotkey(op_id, target, args)

    # ---------- rollback ----------

    def _action_rollback_list(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        max_results = int(args.get("max_results", 50))
        ops = []
        for snap in sorted(self.backup_dir.glob("*/snapshots.json"), reverse=True):
            try:
                snapshots = json.loads(snap.read_text(encoding="utf-8"))
                ops.append({"op_id": snap.parent.name, "snapshot_count": len(snapshots), "paths": [s.get("rel_path") for s in snapshots]})
            except Exception:
                continue
            if len(ops) >= max_results:
                break
        return {"operations": ops}

    def _action_rollback_apply(self, op_id: str, target: Optional[str], args: Dict[str, Any]) -> Dict[str, Any]:
        target_op_id = args.get("op_id") or target
        if not target_op_id:
            raise OmniBodyError("rollback.apply requires target or args.op_id")
        snap_path = self.backup_dir / str(target_op_id) / "snapshots.json"
        if not snap_path.exists():
            raise FileNotFoundError(f"No snapshots found for op_id={target_op_id}")
        snapshots = json.loads(snap_path.read_text(encoding="utf-8"))
        result = self._restore_snapshots(snapshots)
        return {"rolled_back_op_id": target_op_id, "result": result}


if __name__ == "__main__":
    raise SystemExit("Direct BodyRuntime execution is disabled; use the signed omni_body entrypoint.")
