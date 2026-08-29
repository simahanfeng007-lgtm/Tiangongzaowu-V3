"""内容级验收：office 交付物空壳检测（真机 2026-08-29 复现修复）。

复现背景：模型生成"只有标题的docx + 只有一个空单元格的xlsx"，回复中
谎报内容详实（列名/人名/分数俱全），存在性与完整性质检全绿放行。
"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from v3.simple_chain.kernel import _office_content_gaps


def _make_shell_xlsx(path: Path) -> None:
    import openpyxl

    openpyxl.Workbook().save(path)


def _make_filled_xlsx(path: Path) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["姓名", "分数"])
    ws.append(["张三", 87])
    ws.append(["李四", 92])
    ws.append(["王五", 78])
    wb.save(path)


def _make_shell_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("项目周报", level=1)
    doc.add_heading("本周完成", level=2)
    doc.add_heading("下周计划", level=2)
    doc.save(path)


def _make_filled_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("项目周报", level=1)
    doc.add_heading("本周完成", level=2)
    for line in (
        "完成智能体核心对话模块的开发与联调，覆盖桌面与移动端。",
        "完成文件操作链路优化，写入校验与回滚机制全部落地。",
        "完成证据识别机制修复，四类工具结果契约全部纳入验收。",
    ):
        doc.add_paragraph(line)
    doc.add_heading("下周计划", level=2)
    for line in (
        "推进技术内容规划，完成首批选题与脚本初稿。",
        "整理开源文档结构，补齐快速上手与部署章节。",
        "对接首批外包客户，明确交付标准与排期。",
    ):
        doc.add_paragraph(line)
    doc.save(path)


REQUEST = (
    '复杂文件任务：创建"项目周报.docx"（含"本周完成"和"下周计划"两部分各3条内容）'
    '和"数据表.xlsx"（姓名、分数两列3行示例数据），两个文件都要。'
)


def _make_placeholder_xlsx(path: Path) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["name", "score"])
    ws.append(["name", "score"])
    ws.append(["name", "score"])
    wb.save(path)


class OfficeContentGateTests(unittest.TestCase):
    def test_placeholder_rows_are_rejected(self) -> None:
        """真机第二轮复现：xlsx 数据行全是重复占位符（name/score ×3），
        列名/行数检查全过但内容为假。重复行检测必须拦住。"""
        with tempfile.TemporaryDirectory() as raw:
            xlsx_path = Path(raw) / "数据表.xlsx"
            _make_placeholder_xlsx(xlsx_path)
            gaps = _office_content_gaps(
                "创建数据表.xlsx要姓名分数两列3行数据",
                [{"path": str(xlsx_path)}],
            )
            self.assertTrue(any("占位符" in gap for gap in gaps), gaps)

    def test_shell_files_are_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            docx_path = root / "项目周报.docx"
            xlsx_path = root / "数据表.xlsx"
            _make_shell_docx(docx_path)
            _make_shell_xlsx(xlsx_path)
            gaps = _office_content_gaps(
                REQUEST,
                [{"path": str(docx_path)}, {"path": str(xlsx_path)}],
            )
            joined = "\n".join(gaps)
            self.assertIn("项目周报.docx", joined)
            self.assertIn("空壳", joined)
            self.assertIn("数据表.xlsx", joined)
            self.assertIn("空表", joined)

    def test_filled_files_pass(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            docx_path = root / "项目周报.docx"
            xlsx_path = root / "数据表.xlsx"
            _make_filled_docx(docx_path)
            _make_filled_xlsx(xlsx_path)
            gaps = _office_content_gaps(
                REQUEST,
                [{"path": str(docx_path)}, {"path": str(xlsx_path)}],
            )
            self.assertEqual(gaps, [])

    def test_missing_columns_reported(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            xlsx_path = root / "数据表.xlsx"
            _make_filled_xlsx(xlsx_path)
            import openpyxl

            wb = openpyxl.load_workbook(xlsx_path)
            wb.active.cell(row=1, column=1, value="名字")  # 改掉表头让列缺失
            wb.save(xlsx_path)
            gaps = _office_content_gaps(REQUEST, [{"path": str(xlsx_path)}])
            self.assertTrue(any("缺少列" in gap and "姓名" in gap for gap in gaps), gaps)

    def test_no_attachments_no_gaps(self) -> None:
        self.assertEqual(_office_content_gaps(REQUEST, []), [])
        self.assertEqual(_office_content_gaps(REQUEST, None), [])


if __name__ == "__main__":
    unittest.main()
